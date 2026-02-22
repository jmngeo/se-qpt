"""
Phase 4: RFP Document Export Service

Generates comprehensive RFP (Request for Proposal) documents consolidating
all organization-specific data from Phases 1-4.

Includes LLM-enhanced Word/PDF generation with:
- Core Concept narrative
- Module Goals and Contents
"""

import io
import os
import json
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
from sqlalchemy import text
from flask import current_app
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

# python-docx imports for Word document generation
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# OpenAI for LLM generation
from openai import OpenAI


class Phase4RFPService:
    """Service for Phase 4 RFP Document Export"""

    # Styling constants (matching AVIVA export)
    TITLE_FONT = Font(size=16, bold=True)
    SUBTITLE_FONT = Font(size=12, bold=True)
    SECTION_FONT = Font(size=11, bold=True)
    LABEL_FONT = Font(bold=True)
    HEADER_FONT = Font(bold=True, color="FFFFFF")
    HEADER_FILL = PatternFill(start_color="455A64", end_color="455A64", fill_type="solid")
    SUBHEADER_FILL = PatternFill(start_color="607D8B", end_color="607D8B", fill_type="solid")
    LIGHT_FILL = PatternFill(start_color="ECEFF1", end_color="ECEFF1", fill_type="solid")
    GREEN_FILL = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid")
    YELLOW_FILL = PatternFill(start_color="FFF8E1", end_color="FFF8E1", fill_type="solid")
    RED_FILL = PatternFill(start_color="FFEBEE", end_color="FFEBEE", fill_type="solid")
    THIN_BORDER = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    LEVEL_NAMES = {1: 'Knowing', 2: 'Understanding', 4: 'Applying', 6: 'Mastering'}
    LEVEL_DURATION_HOURS = {1: 1, 2: 2, 4: 4, 6: 8}
    COMPETENCY_AREAS = {
        'Core': ['Systems Thinking', 'Systems Modelling and Analysis', 'Lifecycle Consideration', 'Customer / Value Orientation'],
        'Technical': ['Requirements Definition', 'System Architecting', 'Integration, Verification, Validation', 'Operation and Support'],
        'Management': ['Project Management', 'Decision Management', 'Information Management', 'Configuration Management'],
        'Social/Personal': ['Agile Methods', 'Self-Organization', 'Communication', 'Leadership']
    }

    def __init__(self, db_session):
        """Initialize the service with database session"""
        self.db = db_session
        # Initialize OpenAI client for LLM generation
        self.openai_client = None
        self.model = "gpt-4o-mini"

    def _get_openai_client(self):
        """Lazy initialization of OpenAI client"""
        if self.openai_client is None:
            api_key = os.getenv('OPENAI_API_KEY')
            if not api_key:
                current_app.logger.warning("[RFP] No OPENAI_API_KEY found, LLM features disabled")
                return None
            self.openai_client = OpenAI(api_key=api_key)
        return self.openai_client

    # =========================================================================
    # LLM GENERATION METHODS
    # =========================================================================

    def generate_core_concept(self, data: Dict[str, Any]) -> str:
        """
        Generate organization-specific Core Concept narrative using LLM.

        Returns ~200 words (2 paragraphs) describing the SE qualification initiative.
        Falls back to template if LLM fails.
        """
        client = self._get_openai_client()
        if not client:
            return self._get_core_concept_fallback(data)

        org = data.get('organization', {})
        maturity = data.get('maturity', {})
        strategies = data.get('strategies', [])
        roles = data.get('roles', [])
        target_group = data.get('target_group', {})
        pmt = data.get('pmt_context', {})

        # Extract strategy names
        primary_strategy = strategies[0].get('name', 'Needs-Based Qualification') if strategies else 'Needs-Based Qualification'
        secondary_strategies = [s.get('name') for s in strategies[1:]] if len(strategies) > 1 else []

        # Get core role names
        role_names = [r.get('name') for r in roles[:5]]

        prompt = f"""You are writing the Core Concept section for a Systems Engineering Qualification Program RFP document.
This section introduces the organization's SE qualification initiative to potential training service providers.

ORGANIZATION CONTEXT:
- Organization: {org.get('name', 'The Organization')}
- Industry: {pmt.get('industry', 'Engineering')}
- SE Maturity Level: {org.get('maturity_level', 3)}/5 (Score: {org.get('maturity_score', 50):.1f}/100)
- Assessment Pathway: {maturity.get('assessment_pathway', 'Task-based competency assessment')}

QUALIFICATION PROGRAM SCOPE:
- Target Personnel: {target_group.get('size', 100)} employees
- Roles Identified: {len(roles)}
- Core Personnel to be Trained: {', '.join(role_names) if role_names else 'Systems Engineers, System Architects'}

SELECTED STRATEGIES:
- Primary: {primary_strategy}
- Secondary: {', '.join(secondary_strategies) if secondary_strategies else 'None'}

INSTRUCTIONS:
Write a Core Concept section (2 paragraphs, approximately 150-200 words total) that:
1. First paragraph: Introduces the organization's systematic approach to equip personnel with SE capabilities. Mention it will strengthen SE as a profession and promote standardization.
2. Second paragraph: Briefly describe the curriculum design (levels A-D, from basic to advanced). End with: "The core personnel to be trained include:" followed by a bulleted list of the main roles.

Use formal business language appropriate for an RFP document.
Do NOT use emojis or casual language.
Do NOT mention specific tools or software.
"""

        try:
            response = client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an expert technical writer specializing in Systems Engineering training documentation."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.5,
                max_tokens=500,
                timeout=30
            )
            content = response.choices[0].message.content.strip()
            current_app.logger.info(f"[RFP] Generated Core Concept ({len(content)} chars)")
            return content

        except Exception as e:
            current_app.logger.error(f"[RFP] Core Concept generation failed: {e}")
            return self._get_core_concept_fallback(data)

    def _get_core_concept_fallback(self, data: Dict[str, Any]) -> str:
        """Fallback template for Core Concept when LLM fails"""
        org = data.get('organization', {})
        roles = data.get('roles', [])
        target_group = data.get('target_group', {})

        role_names = [r.get('name') for r in roles[:5]]
        role_list = '\n'.join([f"- {name}" for name in role_names]) if role_names else "- Systems Engineers\n- System Architects"

        return f"""{org.get('name', 'The organization')} is planning a systematic approach to equip personnel with the necessary capabilities to excel in their projects. This initiative will strengthen Systems Engineering as a profession and promote standardization across locations. Systems Engineering encompasses multiple aspects, and thus, the curriculum will reflect this diversity. The initial focus will be on interrelated topics including systems architecture, driven by Model-Based Systems Engineering (MBSE) techniques, and traditional requirements engineering.

The curriculum is designed to have levels ranging from A (lowest) to D (highest). Since personnel are not expected to work in isolation, they need to share a specific set of competencies with other engineering disciplines. Therefore, some level A training modules will be available for multiple disciplines.

The core personnel to be trained include:
{role_list}

The emphasis on outsourcing training provision is on the lower levels, as there is a significantly larger number of personnel ({target_group.get('size', 'approximately 100')} employees) that need to be trained."""

    def generate_module_goals(self, module_data: Dict[str, Any]) -> List[str]:
        """
        Generate 3-5 level-appropriate goal bullet points for a training module.

        Uses level-appropriate verbs:
        - Level 1 (Knowing): "Understanding...", "Knowing...", "Being aware of..."
        - Level 2 (Understanding): "Understanding...", "Comprehending...", "Explaining..."
        - Level 4 (Applying): "Applying...", "Being able to...", "Performing..."
        - Level 6 (Mastering): "Mastering...", "Optimizing...", "Teaching..."
        """
        client = self._get_openai_client()
        if not client:
            return self._get_module_goals_fallback(module_data)

        competency_name = module_data.get('competency_name', 'Systems Engineering')
        target_level = module_data.get('target_level', 2)
        level_name = self.LEVEL_NAMES.get(target_level, 'Understanding')
        learning_objective = module_data.get('learning_objective', '')
        content_topics = module_data.get('content_topics', [])
        pmt_type = module_data.get('pmt_type', 'combined')

        # Level-appropriate verb guidance
        level_verbs = {
            1: "Use verbs like: 'Understanding...', 'Knowing...', 'Being aware of...', 'Recognizing...'",
            2: "Use verbs like: 'Understanding...', 'Comprehending...', 'Explaining...', 'Describing...'",
            4: "Use verbs like: 'Applying...', 'Being able to...', 'Performing...', 'Implementing...'",
            6: "Use verbs like: 'Mastering...', 'Optimizing...', 'Teaching...', 'Leading...'"
        }

        prompt = f"""Generate 3-5 learning goals for this Systems Engineering training module.

MODULE: {competency_name} - {level_name}
TARGET LEVEL: {target_level} ({level_name})
PMT FOCUS: {pmt_type}

LEARNING OBJECTIVE (formal):
{learning_objective if learning_objective else 'Enable participants to develop competency in ' + competency_name}

CONTENT TOPICS:
{chr(10).join('- ' + t for t in content_topics[:8]) if content_topics else '- General topics for ' + competency_name}

INSTRUCTIONS:
1. {level_verbs.get(target_level, level_verbs[2])}
2. Each goal should be one line, starting with the verb
3. Cover main learning outcomes for this competency
4. Keep goals concise and actionable
5. Be appropriate for {pmt_type} focus (process/method/tool)

Return ONLY a JSON array of 3-5 goal strings:
["Understanding the role of...", "Understanding how...", ...]
"""

        try:
            response = client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an instructional designer for Systems Engineering training. Output only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=400,
                timeout=30
            )
            content = response.choices[0].message.content.strip()

            # Handle markdown code blocks
            if content.startswith('```'):
                content = content.split('```')[1]
                if content.startswith('json'):
                    content = content[4:]
                content = content.strip()

            goals = json.loads(content)
            if isinstance(goals, list) and len(goals) >= 3:
                return goals[:5]  # Max 5 goals

        except Exception as e:
            current_app.logger.warning(f"[RFP] Module goals generation failed for {competency_name}: {e}")

        return self._get_module_goals_fallback(module_data)

    def _get_module_goals_fallback(self, module_data: Dict[str, Any]) -> List[str]:
        """Fallback goals when LLM fails"""
        competency_name = module_data.get('competency_name', 'the competency')
        target_level = module_data.get('target_level', 2)
        learning_objective = module_data.get('learning_objective', '')

        # Level-appropriate default goals
        if target_level <= 1:
            return [
                f"Understanding the role of {competency_name} in system development",
                f"Knowing the basic concepts and terminology",
                f"Being aware of key principles and guidelines"
            ]
        elif target_level <= 2:
            return [
                f"Understanding the principles of {competency_name}",
                f"Comprehending the relationships to other SE activities",
                f"Explaining key concepts and their applications"
            ]
        elif target_level <= 4:
            return [
                f"Applying {competency_name} techniques in practice",
                f"Being able to perform key tasks independently",
                f"Implementing solutions following established procedures",
                learning_objective if learning_objective else f"Demonstrating competency in {competency_name}"
            ]
        else:
            return [
                f"Mastering advanced {competency_name} techniques",
                f"Optimizing processes and mentoring others",
                f"Leading {competency_name} activities across teams"
            ]

    def generate_module_contents(self, module_data: Dict[str, Any]) -> List[str]:
        """
        Generate 6-10 numbered content items for a training module.

        Uses content_topics from competency_content_baseline as input.
        """
        client = self._get_openai_client()
        if not client:
            return self._get_module_contents_fallback(module_data)

        competency_name = module_data.get('competency_name', 'Systems Engineering')
        target_level = module_data.get('target_level', 2)
        level_name = self.LEVEL_NAMES.get(target_level, 'Understanding')
        content_topics = module_data.get('content_topics', [])
        duration_hours = module_data.get('duration_hours', 8)
        org_tools = module_data.get('org_tools', '')
        learning_format = module_data.get('learning_format', 'Seminar')

        prompt = f"""Generate a numbered content outline for this Systems Engineering training module.

MODULE: {competency_name} - {level_name}
TARGET LEVEL: {target_level} ({level_name})
DURATION: {duration_hours} hours
FORMAT: {learning_format}

CONTENT TOPICS (baseline from curriculum):
{chr(10).join('- ' + t for t in content_topics) if content_topics else '- Core concepts of ' + competency_name}

ORGANIZATION TOOLS: {org_tools if org_tools else 'Standard SE tools'}

INSTRUCTIONS:
1. Organize topics into a logical training sequence (6-10 main items)
2. Include sub-topics where appropriate (a, b, c)
3. Adjust depth based on Level {target_level}:
   - Level 1: Overview, definitions, basic concepts
   - Level 2: Concepts, relationships, examples
   - Level 4: Hands-on application, tool usage, practical exercises
   - Level 6: Advanced techniques, optimization, mentoring
4. Start with "Introduction/Recap" and end with "Summary/Application"
5. Keep realistic for {duration_hours} hours

Return ONLY a JSON array of content strings:
["1. Introduction and Recap", "2. Topic A", "   a. Sub-topic", "3. Topic B", ...]
"""

        try:
            response = client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an instructional designer for Systems Engineering training. Output only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.5,
                max_tokens=600,
                timeout=30
            )
            content = response.choices[0].message.content.strip()

            # Handle markdown code blocks
            if content.startswith('```'):
                content = content.split('```')[1]
                if content.startswith('json'):
                    content = content[4:]
                content = content.strip()

            contents = json.loads(content)
            if isinstance(contents, list) and len(contents) >= 5:
                return contents[:12]  # Max 12 items

        except Exception as e:
            current_app.logger.warning(f"[RFP] Module contents generation failed for {competency_name}: {e}")

        return self._get_module_contents_fallback(module_data)

    def _get_module_contents_fallback(self, module_data: Dict[str, Any]) -> List[str]:
        """Fallback contents when LLM fails - use content_topics directly"""
        content_topics = module_data.get('content_topics', [])
        competency_name = module_data.get('competency_name', 'the topic')

        if content_topics:
            contents = ["1. Introduction and Recap"]
            for i, topic in enumerate(content_topics[:8], start=2):
                contents.append(f"{i}. {topic}")
            contents.append(f"{len(contents) + 1}. Summary and Application")
            return contents
        else:
            return [
                "1. Introduction and Recap",
                f"2. Fundamentals of {competency_name}",
                "3. Key Concepts and Terminology",
                "4. Methods and Techniques",
                "5. Practical Application",
                "6. Best Practices",
                "7. Summary and Q&A"
            ]

    def generate_service_requirements(self, data: Dict[str, Any]) -> List[str]:
        """
        Generate service requirements/constraints bullet points.

        Conditional based on:
        - Format selection (seminar=travel, e-learning=no travel)
        - PMT tools available
        - Timeline constraints
        """
        org = data.get('organization', {})
        pmt = data.get('pmt_context', {})
        phase3 = data.get('phase3', {})
        timeline = phase3.get('timeline', {})
        config = phase3.get('config', {})

        # Determine format distribution
        format_dist = phase3.get('summary', {}).get('format_distribution', {})
        has_seminar = format_dist.get('Seminar', 0) > 0 or format_dist.get('Workshop', 0) > 0
        has_elearning = format_dist.get('E-Learning', 0) > 0 or format_dist.get('Self-Study', 0) > 0

        requirements = []

        # 1. Trainer proficiency (always)
        requirements.append("Trainer shall be proficient in systems engineering, especially in the topics to be trained as well as in providing professional trainings.")

        # 2. PMT processes (if available and high maturity)
        if pmt.get('processes') and org.get('maturity_level', 0) >= 4:
            requirements.append(f"For some training modules the trainer shall be able to delve into {org.get('name', 'organization')} processes.")

        # 3. Travel (conditional on format)
        if has_seminar and not has_elearning:
            requirements.append("Trainer shall be willing to travel to respective organization locations. The trainings will be on-site.")
        elif has_elearning and not has_seminar:
            requirements.append("Training delivery will be primarily online/remote. No travel is required.")
        elif has_seminar and has_elearning:
            requirements.append("Trainer shall be willing to travel for on-site sessions. Some modules may be delivered remotely.")

        # 4. Language (always)
        requirements.append("Trainer shall be able to provide the training in English.")

        # 5. Tools (if specified)
        if pmt.get('tools'):
            tools = pmt.get('tools')
            requirements.append(f"The trainings shall include hands-on exercises with the given tools: {tools}")

        # 6. Group size (from format max_participants)
        max_participants = 12  # Default
        modules = phase3.get('modules', [])
        if modules:
            selected_formats = [m.get('selected_format', {}) for m in modules if m.get('selected_format')]
            if selected_formats:
                max_vals = [f.get('max_participants', 12) for f in selected_formats if f.get('max_participants')]
                if max_vals:
                    max_participants = max(max_vals)
        requirements.append(f"The size of the training groups will not exceed {max_participants} participants.")

        # 7. Timeline (if available)
        milestones = timeline.get('milestones', [])
        if milestones and len(milestones) >= 2:
            first = milestones[0].get('estimated_date', '')
            last = milestones[-1].get('estimated_date', '')
            if first and last:
                requirements.append(f"All personnel shall be trained within the timeline from {first} to {last}.")

        return requirements

    def _generate_single_module_content_with_app(
        self,
        module_data: Dict[str, Any],
        app
    ) -> Tuple[int, List[str], List[str]]:
        """
        Generate goals and contents for a single module within Flask app context.
        Used by parallel processing to generate multiple modules concurrently.

        Args:
            module_data: Module data dict with _module_index
            app: Flask app instance for context

        Returns: (module_index, goals_list, contents_list)
        """
        module_idx = module_data.get('_module_index', 0)
        with app.app_context():
            try:
                goals = self.generate_module_goals(module_data)
                contents = self.generate_module_contents(module_data)
                return (module_idx, goals, contents)
            except Exception as e:
                # Log error and return fallback
                print(f"[RFP] Parallel generation failed for module {module_idx}: {e}")
                return (
                    module_idx,
                    self._get_module_goals_fallback(module_data),
                    self._get_module_contents_fallback(module_data)
                )

    def generate_all_module_content_parallel(
        self,
        modules: List[Dict[str, Any]],
        pmt_context: Dict[str, Any],
        max_workers: int = 5
    ) -> Dict[int, Dict[str, Any]]:
        """
        Generate LLM content for all modules in parallel.

        This significantly speeds up Word/PDF export by processing multiple
        modules concurrently instead of sequentially.

        Args:
            modules: List of module data from Phase 3
            pmt_context: Organization PMT context
            max_workers: Maximum parallel threads (default 5 to respect API limits)

        Returns:
            Dict mapping module index to generated content:
            {0: {'goals': [...], 'contents': [...]}, 1: {...}, ...}
        """
        # Capture Flask app reference for thread context
        app = current_app._get_current_object()
        current_app.logger.info(f"[RFP] Starting parallel LLM generation for {len(modules)} modules with {max_workers} workers")

        # Prepare module data for parallel processing
        prepared_modules = []
        for idx, module in enumerate(modules):
            competency_id = module.get('competency_id')
            content_topics = self.get_content_topics_for_competency(competency_id) if competency_id else []

            selected_format = module.get('selected_format') or {}
            format_name = (module.get('learning_format_name') or
                          selected_format.get('format_name') or 'N/A')

            module_data = {
                '_module_index': idx,
                'competency_name': module.get('competency_name', 'Unknown'),
                'target_level': module.get('target_level', 2),
                'learning_objective': module.get('learning_objective', ''),
                'content_topics': content_topics,
                'pmt_type': module.get('pmt_type', 'combined'),
                'duration_hours': module.get('estimated_duration_hours', 8),
                'org_tools': pmt_context.get('tools', ''),
                'learning_format': format_name
            }
            prepared_modules.append(module_data)

        # Generate content in parallel with app context
        results = {}
        completed_count = 0
        total_count = len(modules)

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all tasks with app context
            future_to_module = {
                executor.submit(
                    self._generate_single_module_content_with_app,
                    mod_data,
                    app
                ): mod_data['_module_index']
                for mod_data in prepared_modules
            }

            # Collect results as they complete
            for future in as_completed(future_to_module):
                try:
                    module_idx, goals, contents = future.result(timeout=60)
                    results[module_idx] = {
                        'goals': goals,
                        'contents': contents
                    }
                    completed_count += 1
                    print(f"[RFP] Completed module {completed_count}/{total_count}")
                except Exception as e:
                    module_idx = future_to_module[future]
                    print(f"[RFP] Module {module_idx} generation error: {e}")
                    # Use fallback for failed modules
                    results[module_idx] = {
                        'goals': self._get_module_goals_fallback(prepared_modules[module_idx]),
                        'contents': self._get_module_contents_fallback(prepared_modules[module_idx])
                    }
                    completed_count += 1

        current_app.logger.info(f"[RFP] Parallel generation complete. Generated {len(results)} modules.")
        return results

    def get_content_topics_for_competency(self, competency_id: int) -> List[str]:
        """Get content topics from competency_content_baseline table"""
        result = self.db.execute(
            text("""
                SELECT content_topics
                FROM competency_content_baseline
                WHERE competency_id = :comp_id
            """),
            {'comp_id': competency_id}
        ).fetchone()

        if result and result.content_topics:
            return result.content_topics if isinstance(result.content_topics, list) else []
        return []

    # =========================================================================
    # DATA AGGREGATION - PHASE 1
    # =========================================================================

    def get_organization_profile(self, organization_id: int) -> Dict[str, Any]:
        """Get organization profile including maturity data"""
        result = self.db.execute(
            text("""
                SELECT
                    o.id,
                    o.organization_name,
                    o.organization_public_key,
                    o.maturity_score,
                    o.selected_archetype,
                    o.phase1_completed,
                    o.phase2_completed,
                    o.phase3_completed,
                    o.created_at
                FROM organization o
                WHERE o.id = :org_id
            """),
            {'org_id': organization_id}
        ).fetchone()

        if not result:
            return {}

        # Determine maturity level from score
        score = result.maturity_score or 0
        if score >= 80:
            maturity_level = 5
        elif score >= 60:
            maturity_level = 4
        elif score >= 40:
            maturity_level = 3
        elif score >= 20:
            maturity_level = 2
        else:
            maturity_level = 1

        return {
            'id': result.id,
            'name': result.organization_name,
            'public_key': result.organization_public_key,
            'maturity_score': score,
            'maturity_level': maturity_level,
            'selected_archetype': result.selected_archetype,
            'phase1_completed': result.phase1_completed,
            'phase2_completed': result.phase2_completed,
            'phase3_completed': result.phase3_completed,
            'created_at': result.created_at.strftime('%Y-%m-%d') if result.created_at else None
        }

    def get_maturity_assessment(self, organization_id: int) -> Dict[str, Any]:
        """Get 4 Fields of Action maturity assessment details"""
        result = self.db.execute(
            text("""
                SELECT
                    questionnaire_type,
                    responses,
                    completed_at
                FROM phase_questionnaire_responses
                WHERE organization_id = :org_id
                  AND questionnaire_type = 'maturity'
                ORDER BY completed_at DESC
                LIMIT 1
            """),
            {'org_id': organization_id}
        ).fetchone()

        if not result:
            return {'fields_of_action': {}, 'overall_score': 0}

        responses = result.responses or {}
        if isinstance(responses, str):
            responses = json.loads(responses)

        # Extract data from the correct JSON structure
        answers = responses.get('answers', {})
        results_data = responses.get('results', {})
        field_scores = results_data.get('fieldScores', {})

        # Field mapping with max scores (SE Roles has 6 levels 0-5, others have 5 levels 0-4)
        field_config = {
            'rolloutScope': {'label': 'Rollout Scope', 'max_score': 5},
            'seRolesProcesses': {'label': 'SE Roles & Processes', 'max_score': 6},
            'seMindset': {'label': 'SE Mindset', 'max_score': 5},
            'knowledgeBase': {'label': 'Knowledge Base', 'max_score': 5}
        }

        # Maturity level descriptions for display
        level_descriptions = {
            'rolloutScope': {
                0: 'Not Available', 1: 'Ad Hoc/Undefined', 2: 'Individually Controlled',
                3: 'Defined and Established', 4: 'Continuously Optimized'
            },
            'seRolesProcesses': {
                0: 'Not Available', 1: 'Ad Hoc/Undefined', 2: 'Individually Controlled',
                3: 'Defined and Established', 4: 'Quantitatively Managed', 5: 'Continuously Optimized'
            },
            'seMindset': {
                0: 'Not Available', 1: 'Ad Hoc/Undefined', 2: 'Individually Controlled',
                3: 'Defined and Established', 4: 'Continuously Optimized'
            },
            'knowledgeBase': {
                0: 'Not Available', 1: 'Ad Hoc/Undefined', 2: 'Individually Controlled',
                3: 'Defined and Established', 4: 'Continuously Optimized'
            }
        }

        fields_of_action = {}
        for key, config in field_config.items():
            answer_value = answers.get(key, 0)
            field_score_percent = field_scores.get(key, 0)
            description = level_descriptions.get(key, {}).get(answer_value, '')

            fields_of_action[config['label']] = {
                'answer_value': answer_value,
                'max_value': config['max_score'],
                'score_percent': field_score_percent,
                'description': description
            }

        # Determine assessment pathway based on SE Roles answer
        se_roles_answer = answers.get('seRolesProcesses', 0)
        assessment_pathway = 'Role-based competency assessment' if se_roles_answer >= 4 else 'Task-based competency assessment'

        overall_score = results_data.get('finalScore', results_data.get('rawScore', 0))

        return {
            'fields_of_action': fields_of_action,
            'overall_score': overall_score,
            'assessment_pathway': assessment_pathway,
            'maturity_level': results_data.get('maturityLevel', 0),
            'maturity_name': results_data.get('maturityName', ''),
            'completed_at': result.completed_at.strftime('%Y-%m-%d') if result.completed_at else None
        }

    def get_target_group_size(self, organization_id: int) -> Dict[str, Any]:
        """Get target group size from Phase 1"""
        result = self.db.execute(
            text("""
                SELECT
                    (responses::jsonb)->>'value' as target_size,
                    (responses::jsonb)->>'range' as range_label
                FROM phase_questionnaire_responses
                WHERE organization_id = :org_id
                  AND questionnaire_type = 'target_group'
                ORDER BY completed_at DESC
                LIMIT 1
            """),
            {'org_id': organization_id}
        ).fetchone()

        return {
            'size': int(result.target_size) if result and result.target_size else 0,
            'range_label': result.range_label if result else 'Unknown'
        }

    def get_strategies(self, organization_id: int) -> List[Dict[str, Any]]:
        """Get selected qualification strategies"""
        results = self.db.execute(
            text("""
                SELECT
                    ls.id,
                    ls.strategy_template_id,
                    st.strategy_name,
                    st.strategy_description,
                    ls.priority,
                    ls.selected
                FROM learning_strategy ls
                JOIN strategy_template st ON ls.strategy_template_id = st.id
                WHERE ls.organization_id = :org_id
                  AND ls.selected = true
                ORDER BY ls.priority ASC
            """),
            {'org_id': organization_id}
        ).fetchall()

        # Use richer descriptions from strategy_selection_engine as fallback
        from app.strategy_selection_engine import SE_TRAINING_STRATEGIES
        strategy_detail_map = {s['name']: s for s in SE_TRAINING_STRATEGIES}

        strategies = []
        for r in results:
            db_desc = r.strategy_description or ''
            rich_desc = strategy_detail_map.get(r.strategy_name, {}).get('description', '')
            strategies.append({
                'id': r.strategy_template_id,
                'name': r.strategy_name,
                'description': rich_desc or db_desc,
                'priority': r.priority,
                'is_primary': r.priority == 1
            })

        return strategies

    def get_organization_roles(self, organization_id: int) -> List[Dict[str, Any]]:
        """Get organization roles with cluster mappings"""
        results = self.db.execute(
            text("""
                SELECT
                    r.id,
                    r.role_name,
                    r.role_description,
                    r.identification_method,
                    r.standard_role_cluster_id,
                    rc.role_cluster_name as se_cluster_name,
                    r.training_program_cluster_id,
                    tpc.training_program_name
                FROM organization_roles r
                LEFT JOIN role_cluster rc ON r.standard_role_cluster_id = rc.id
                LEFT JOIN training_program_cluster tpc ON r.training_program_cluster_id = tpc.id
                WHERE r.organization_id = :org_id
                ORDER BY r.id
            """),
            {'org_id': organization_id}
        ).fetchall()

        roles = []
        for r in results:
            roles.append({
                'id': r.id,
                'name': r.role_name,
                'description': r.role_description,
                'identification_method': r.identification_method,
                'se_cluster_id': r.standard_role_cluster_id,
                'se_cluster_name': r.se_cluster_name,
                'training_program_id': r.training_program_cluster_id,
                'training_program': r.training_program_name
            })

        return roles

    # =========================================================================
    # DATA AGGREGATION - PHASE 2
    # =========================================================================

    def get_competency_gaps(self, organization_id: int) -> Dict[str, Any]:
        """Get aggregated competency gap analysis from generated learning objectives"""
        # Get gap data from generated_learning_objectives (computed in Phase 2)
        result = self.db.execute(
            text("""
                SELECT objectives_data
                FROM generated_learning_objectives
                WHERE organization_id = :org_id
                ORDER BY generated_at DESC
                LIMIT 1
            """),
            {'org_id': organization_id}
        ).fetchone()

        gaps_by_competency = {}
        gaps_by_area = {'Core': 0, 'Technical': 0, 'Management': 0, 'Social/Personal': 0}
        gaps_by_level = {1: 0, 2: 0, 4: 0, 6: 0}
        total_gaps = 0

        if not result or not result.objectives_data:
            return {
                'total_gaps': 0,
                'by_competency': gaps_by_competency,
                'by_area': gaps_by_area,
                'by_level': gaps_by_level
            }

        data = result.objectives_data
        if isinstance(data, str):
            data = json.loads(data)

        # Extract gap data from the nested pyramid structure
        main_pyramid = data.get('data', {}).get('main_pyramid', {})
        levels_data = main_pyramid.get('levels', {})

        # Map competency IDs to areas
        competency_areas = {}
        for comp in self.COMPETENCY_AREAS.get('Core', []):
            competency_areas[comp] = 'Core'
        for comp in self.COMPETENCY_AREAS.get('Technical', []):
            competency_areas[comp] = 'Technical'
        for comp in self.COMPETENCY_AREAS.get('Management', []):
            competency_areas[comp] = 'Management'
        for comp in self.COMPETENCY_AREAS.get('Social/Personal', []):
            competency_areas[comp] = 'Social/Personal'

        # Process each level's competencies
        seen_competencies = set()
        for level_key, level_data in levels_data.items():
            competencies = level_data.get('competencies', [])
            for comp in competencies:
                gap_data = comp.get('gap_data')
                if not gap_data or not gap_data.get('has_gap'):
                    continue

                comp_id = comp.get('competency_id')
                comp_name = comp.get('competency_name', f'Competency {comp_id}')
                levels_needed = gap_data.get('levels_needed', [])

                if not levels_needed:
                    continue

                # Only count each competency once (not per level)
                if comp_id not in seen_competencies:
                    seen_competencies.add(comp_id)

                    # Determine area
                    area = competency_areas.get(comp_name, 'Core')

                    gaps_by_competency[comp_name] = {
                        'competency_id': comp_id,
                        'area': area,
                        'levels_needed': levels_needed,
                        'gap_data': gap_data
                    }

                    # Count by area
                    if area in gaps_by_area:
                        gaps_by_area[area] += 1

                    # Count by level
                    for level in levels_needed:
                        if level in gaps_by_level:
                            gaps_by_level[level] += 1

                    total_gaps += 1

        return {
            'total_gaps': total_gaps,
            'by_competency': gaps_by_competency,
            'by_area': gaps_by_area,
            'by_level': gaps_by_level
        }

    def get_existing_trainings(self, organization_id: int) -> List[Dict[str, Any]]:
        """Get existing training coverage"""
        results = self.db.execute(
            text("""
                SELECT
                    et.id,
                    et.competency_id,
                    c.competency_name,
                    et.covered_levels,
                    et.notes
                FROM organization_existing_trainings et
                JOIN competency c ON et.competency_id = c.id
                WHERE et.organization_id = :org_id
                ORDER BY c.competency_name
            """),
            {'org_id': organization_id}
        ).fetchall()

        trainings = []
        for r in results:
            covered = r.covered_levels or []
            if isinstance(covered, str):
                covered = json.loads(covered)

            trainings.append({
                'competency_id': r.competency_id,
                'competency_name': r.competency_name,
                'covered_levels': covered,
                'notes': r.notes
            })

        return trainings

    def get_pmt_context(self, organization_id: int) -> Dict[str, Any]:
        """Get organization's PMT context"""
        result = self.db.execute(
            text("""
                SELECT
                    processes,
                    methods,
                    tools,
                    industry,
                    additional_context
                FROM organization_pmt_context
                WHERE organization_id = :org_id
            """),
            {'org_id': organization_id}
        ).fetchone()

        if not result:
            return {'processes': '', 'methods': '', 'tools': '', 'industry': '', 'additional_context': ''}

        # PMT context fields are stored as text descriptions, not arrays
        return {
            'processes': result.processes or '',
            'methods': result.methods or '',
            'tools': result.tools or '',
            'industry': result.industry or '',
            'additional_context': result.additional_context or ''
        }

    def get_learning_objectives(self, organization_id: int) -> Dict[str, Any]:
        """Get generated learning objectives from Phase 2"""
        result = self.db.execute(
            text("""
                SELECT objectives_data, generated_at
                FROM generated_learning_objectives
                WHERE organization_id = :org_id
                ORDER BY generated_at DESC
                LIMIT 1
            """),
            {'org_id': organization_id}
        ).fetchone()

        if not result:
            return {}

        data = result.objectives_data
        if isinstance(data, str):
            data = json.loads(data)

        return {
            'data': data,
            'generated_at': result.generated_at.strftime('%Y-%m-%d %H:%M') if result.generated_at else None
        }

    # =========================================================================
    # DATA AGGREGATION - PHASE 3
    # =========================================================================

    def get_phase3_data(self, organization_id: int) -> Dict[str, Any]:
        """Get Phase 3 data using existing service"""
        from app.services.phase3_planning_service import Phase3PlanningService
        phase3_service = Phase3PlanningService(self.db)
        return phase3_service.get_phase3_output(organization_id)

    def get_role_competency_matrix(self, organization_id: int) -> List[Dict[str, Any]]:
        """Get role-competency matrix"""
        results = self.db.execute(
            text("""
                SELECT
                    r.id as role_id,
                    r.role_name,
                    c.id as competency_id,
                    c.competency_name,
                    rcm.role_competency_value as required_level
                FROM role_competency_matrix rcm
                JOIN organization_roles r ON rcm.role_cluster_id = r.id
                JOIN competency c ON rcm.competency_id = c.id
                WHERE rcm.organization_id = :org_id
                  AND rcm.role_competency_value > 0
                ORDER BY r.role_name, c.id
            """),
            {'org_id': organization_id}
        ).fetchall()

        matrix = []
        for r in results:
            matrix.append({
                'role_id': r.role_id,
                'role_name': r.role_name,
                'competency_id': r.competency_id,
                'competency_name': r.competency_name,
                'required_level': r.required_level
            })

        return matrix

    # =========================================================================
    # DATA AGGREGATION - PHASE 4 (AVIVA)
    # =========================================================================

    def get_aviva_plans(self, organization_id: int) -> List[Dict[str, Any]]:
        """Get all AVIVA plans for the organization"""
        from app.services.phase4_aviva_service import Phase4AvivaService
        aviva_service = Phase4AvivaService(self.db)

        # Get modules with AVIVA plans
        result = aviva_service.get_modules_for_aviva(organization_id)
        modules = result.get('modules', [])

        plans = []
        for m in modules:
            if m.get('has_aviva_plan') and m.get('id'):
                plan = aviva_service.get_aviva_plan(m['id'])
                if plan:
                    # Merge module info
                    plan['estimated_participants'] = m.get('estimated_participants', 0)
                    plan['roles_needing_training'] = m.get('roles_needing_training', [])
                    plan['subcluster'] = m.get('subcluster')
                    if not plan.get('cluster_name'):
                        plan['cluster_name'] = m.get('cluster_name')
                    plans.append(plan)

        return plans

    # =========================================================================
    # COMPLETE RFP DATA AGGREGATION
    # =========================================================================

    def get_rfp_data(self, organization_id: int, include_aviva: bool = False,
                     module_ids: Optional[List[int]] = None) -> Dict[str, Any]:
        """Aggregate all data needed for RFP export.

        Args:
            organization_id: Organization ID
            include_aviva: Whether to include AVIVA plan data
            module_ids: Optional list of module IDs to filter (from AVIVA selection).
                       If provided, only these modules are included in the export.
        """
        current_app.logger.info(f"[RFP] Aggregating data for organization {organization_id}"
                                f", module_ids filter: {len(module_ids) if module_ids else 'all'}")

        # Phase 1 data
        org_profile = self.get_organization_profile(organization_id)
        maturity = self.get_maturity_assessment(organization_id)
        target_group = self.get_target_group_size(organization_id)
        strategies = self.get_strategies(organization_id)
        roles = self.get_organization_roles(organization_id)

        # Phase 2 data
        gaps = self.get_competency_gaps(organization_id)
        existing_trainings = self.get_existing_trainings(organization_id)
        pmt_context = self.get_pmt_context(organization_id)
        learning_objectives = self.get_learning_objectives(organization_id)

        # Phase 3 data
        phase3_data = self.get_phase3_data(organization_id)

        # Filter modules by selected IDs if provided
        if module_ids is not None and phase3_data.get('modules'):
            module_ids_set = set(module_ids)
            all_count = len(phase3_data['modules'])
            phase3_data['modules'] = [
                m for m in phase3_data['modules'] if m.get('id') in module_ids_set
            ]
            filtered_count = len(phase3_data['modules'])
            current_app.logger.info(f"[RFP] Filtered modules: {filtered_count}/{all_count} (selected from AVIVA)")

        # Role-competency matrix (if role-based)
        role_competency_matrix = []
        if maturity.get('assessment_pathway') == 'Role-based competency assessment':
            role_competency_matrix = self.get_role_competency_matrix(organization_id)

        return {
            'organization': org_profile,
            'maturity': maturity,
            'target_group': target_group,
            'strategies': strategies,
            'roles': roles,
            'gaps': gaps,
            'existing_trainings': existing_trainings,
            'pmt_context': pmt_context,
            'learning_objectives': learning_objectives,
            'phase3': phase3_data,
            'role_competency_matrix': role_competency_matrix,
            'export_timestamp': datetime.now().strftime('%Y-%m-%d %H:%M')
        }

    # =========================================================================
    # EXCEL EXPORT
    # =========================================================================

    def export_rfp_to_excel(self, organization_id: int, include_aviva: bool = False,
                            module_ids: Optional[List[int]] = None) -> io.BytesIO:
        """
        Export comprehensive RFP document to Excel.

        Creates a multi-sheet workbook with all Phase 1-3 data.

        Args:
            organization_id: Organization ID
            include_aviva: Whether to include AVIVA data
            module_ids: Optional list of module IDs to filter (from AVIVA selection)
        """
        current_app.logger.info(f"[RFP] Starting Excel export for organization {organization_id}")

        # Get all data (with optional module filtering)
        data = self.get_rfp_data(organization_id, include_aviva=False, module_ids=module_ids)

        wb = Workbook()

        # Sheet 1: Summary
        self._write_summary_sheet(wb, data)

        # Sheet 2: Maturity Assessment
        self._write_maturity_sheet(wb, data)

        # Sheet 3: Organization Roles
        self._write_roles_sheet(wb, data)

        # Sheet 4: Training Modules
        self._write_modules_sheet(wb, data)

        # Sheet 5: Training Schedule (replaces Timeline)
        self._write_training_schedule_sheet(wb, data)

        # Sheet 6: Role-Competency Matrix (if available)
        if data.get('role_competency_matrix'):
            self._write_role_competency_sheet(wb, data)

        # Save to buffer
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        current_app.logger.info(f"[RFP] Excel export completed")
        return buffer

    def _write_summary_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """Write the Summary sheet"""
        ws = wb.active
        ws.title = "Summary"

        org = data.get('organization', {})
        maturity = data.get('maturity', {})
        target = data.get('target_group', {})
        strategies = data.get('strategies', [])
        phase3 = data.get('phase3', {})
        summary = phase3.get('summary', {})
        gaps = data.get('gaps', {})
        modules = phase3.get('modules', [])

        # Style for description/annotation rows
        DESC_FONT = Font(italic=True, size=9, color='606266')
        PHASE_FONT = Font(italic=True, size=9, color='909399')

        # Helper: write an overview item with optional description and phase origin
        def write_overview_item(r, label, value, phase_origin=None, description=None):
            ws[f'A{r}'] = label
            ws[f'A{r}'].font = self.LABEL_FONT
            ws[f'B{r}'] = value
            ws.merge_cells(f'B{r}:F{r}')
            if phase_origin:
                ws[f'G{r}'] = phase_origin
                ws[f'G{r}'].font = PHASE_FONT
            r += 1
            if description:
                ws[f'B{r}'] = description
                ws[f'B{r}'].font = DESC_FONT
                ws.merge_cells(f'B{r}:H{r}')
                r += 1
            return r

        row = 1

        # Section header styling for this sheet
        SECTION_FILL = PatternFill(start_color='D6E4F0', end_color='D6E4F0', fill_type='solid')
        SECTION_HDR_FONT = Font(size=11, bold=True, color='2F5496')
        SECTION_BORDER = Border(bottom=Side(style='medium', color='4472C4'))

        def write_section_header(r, title):
            ws.merge_cells(f'A{r}:H{r}')
            cell = ws[f'A{r}']
            cell.value = title
            cell.font = SECTION_HDR_FONT
            cell.fill = SECTION_FILL
            cell.border = SECTION_BORDER
            cell.alignment = Alignment(vertical='center')
            ws.row_dimensions[r].height = 24
            return r + 1

        # Title
        ws.merge_cells('A1:H1')
        ws['A1'] = 'SE Qualification Program - Request for Proposal'
        ws['A1'].font = Font(size=16, bold=True, color='FFFFFF')
        ws['A1'].fill = self.HEADER_FILL
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 36

        # Subtitle
        ws.merge_cells('A2:H2')
        ws['A2'] = f"Organization: {org.get('name', 'Unknown')}  |  Generated: {data.get('export_timestamp', '')}"
        ws['A2'].font = Font(size=10, italic=True, color='606266')
        ws['A2'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[2].height = 22
        row = 4

        # ===== ORGANIZATION PROFILE =====
        row = write_section_header(row, 'ORGANIZATION PROFILE')

        row = write_overview_item(row, 'Organization:', org.get('name', 'Unknown'),
            description='The organization for which this SE qualification program RFP was created.')

        row = write_overview_item(row, 'SE Maturity Level:',
            f"Level {org.get('maturity_level', 1)}/5 (Score: {org.get('maturity_score', 0):.1f}/100)",
            phase_origin='(Phase 1, Task 3)',
            description='Organization SE maturity assessed via the Maturity Assessment questionnaire.')

        row = write_overview_item(row, 'Assessment Pathway:',
            maturity.get('assessment_pathway', 'Task-based competency assessment'),
            phase_origin='(Phase 2)',
            description='The competency assessment method used to identify qualification gaps.')

        row += 1

        # ===== QUALIFICATION STRATEGY =====
        row = write_section_header(row, 'QUALIFICATION STRATEGY')

        # Get strategy descriptions
        from app.strategy_selection_engine import SE_TRAINING_STRATEGIES
        strategy_detail_map = {s['name']: s for s in SE_TRAINING_STRATEGIES}

        if strategies:
            # List all strategies with primary/secondary labels
            labeled = []
            for s in strategies:
                lbl = f"{s['name']} (Primary)" if s.get('is_primary') else s['name']
                labeled.append(lbl)
            row = write_overview_item(row, 'Selected Strategies:', ', '.join(labeled),
                phase_origin='(Phase 1, Task 2)',
                description='Qualification strategies selected to guide the training program design.')

            # Individual strategy descriptions
            for s in strategies:
                s_name = s.get('name', '')
                # Use description from data (DB) first, fall back to engine definitions
                desc = s.get('description', '') or strategy_detail_map.get(s_name, {}).get('description', '')
                detail = strategy_detail_map.get(s_name, {})
                if desc:
                    ws[f'B{row}'] = f"{s_name}{' (Primary)' if s.get('is_primary') else ''}:"
                    ws[f'B{row}'].font = Font(bold=True, size=9)
                    row += 1
                    ws[f'B{row}'] = desc
                    ws[f'B{row}'].font = DESC_FONT
                    ws.merge_cells(f'B{row}:H{row}')
                    ws[f'B{row}'].alignment = Alignment(wrap_text=True)
                    row += 1
                    details_parts = []
                    if detail.get('targetAudience'):
                        details_parts.append(f"Target: {detail['targetAudience']}")
                    if detail.get('qualificationLevel'):
                        details_parts.append(f"Level: {detail['qualificationLevel']}")
                    if detail.get('duration'):
                        details_parts.append(f"Duration: {detail['duration']}")
                    if details_parts:
                        ws[f'B{row}'] = ' | '.join(details_parts)
                        ws[f'B{row}'].font = PHASE_FONT
                        ws.merge_cells(f'B{row}:H{row}')
                        row += 1
        else:
            row = write_overview_item(row, 'Strategy:', 'Not Selected')

        row += 1

        # ===== PROGRAM SCOPE =====
        row = write_section_header(row, 'PROGRAM SCOPE')

        row = write_overview_item(row, 'Target Group Size:',
            f"{target.get('size', 0)} employees ({target.get('range_label', '')})",
            phase_origin='(Phase 1, Task 1)',
            description='Number of employees in the target group to be qualified in SE competencies.')

        # Training View
        view_type = summary.get('view_type', phase3.get('view_type', ''))
        if not view_type:
            # Detect from modules
            view_type = 'role_clustered' if any(m.get('cluster_name') for m in modules) else 'competency_level'
        is_role_clustered = view_type == 'role_clustered'
        view_label = 'Role-Clustered' if is_role_clustered else 'Competency-Level'
        if is_role_clustered:
            view_desc = 'Modules grouped into 3 training packages (SE for Engineers, SE for Managers, SE for Interfacing Partners) based on role cluster assignments.'
        else:
            view_desc = 'Modules grouped by the 16 SE competency areas, each containing sub-modules at identified competency levels.'
        row = write_overview_item(row, 'Training View:', view_label,
            phase_origin='(Phase 3, Task 1)',
            description=view_desc)

        row = write_overview_item(row, 'Total Training Modules:', summary.get('total_modules', len(modules)),
            phase_origin='(Phase 3, Task 2)',
            description='Training modules identified from competency gap analysis and learning format assignment.')

        row = write_overview_item(row, 'Roles Defined:', len(data.get('roles', [])),
            phase_origin='(Phase 2, Task 1)',
            description='Organization roles assessed for SE competency gaps.')

        # Total Estimated Duration
        total_duration = sum(self.LEVEL_DURATION_HOURS.get(m.get('target_level', 0), 2) for m in modules)
        duration_display = f"{total_duration} hours"
        if total_duration >= 8:
            duration_display += f" (~{total_duration / 8:.1f} training days)"
        row = write_overview_item(row, 'Total Est. Duration:', duration_display,
            phase_origin='(Derived)',
            description='Cumulative estimated training duration across all modules (L1=1h, L2=2h, L4=4h, L6=8h per module).')

        # Format distribution
        format_dist = summary.get('format_distribution', {})
        if format_dist:
            format_str = ', '.join(f"{k}: {v}" for k, v in format_dist.items())
            row = write_overview_item(row, 'Format Distribution:', format_str,
                phase_origin='(Phase 3, Task 2)',
                description='Distribution of learning formats selected for training delivery.')

        row += 1

        # PMT Context Section (text descriptions, not arrays)
        pmt = data.get('pmt_context', {})
        if any([pmt.get('tools'), pmt.get('methods'), pmt.get('processes')]):
            row = write_section_header(row, 'PMT CONTEXT')

            if pmt.get('industry'):
                ws[f'A{row}'] = 'Industry:'
                ws[f'A{row}'].font = self.LABEL_FONT
                ws.merge_cells(f'B{row}:H{row}')
                ws[f'B{row}'] = pmt['industry']
                row += 1

            if pmt.get('tools'):
                ws[f'A{row}'] = 'Tools:'
                ws[f'A{row}'].font = self.LABEL_FONT
                ws.merge_cells(f'B{row}:H{row}')
                ws[f'B{row}'] = pmt['tools']
                ws[f'B{row}'].alignment = Alignment(wrap_text=True, vertical='top')
                row += 1

            if pmt.get('methods'):
                ws[f'A{row}'] = 'Methods:'
                ws[f'A{row}'].font = self.LABEL_FONT
                ws.merge_cells(f'B{row}:H{row}')
                ws[f'B{row}'] = pmt['methods']
                ws[f'B{row}'].alignment = Alignment(wrap_text=True, vertical='top')
                row += 1

            if pmt.get('processes'):
                ws[f'A{row}'] = 'Processes:'
                ws[f'A{row}'].font = self.LABEL_FONT
                ws.merge_cells(f'B{row}:H{row}')
                ws[f'B{row}'] = pmt['processes']
                ws[f'B{row}'].alignment = Alignment(wrap_text=True, vertical='top')
                row += 1

        # Set column widths (B:H are merged for values, so B just needs moderate width)
        ws.column_dimensions['A'].width = 22
        ws.column_dimensions['B'].width = 18
        ws.column_dimensions['G'].width = 18
        ws.column_dimensions['H'].width = 14

    def _write_maturity_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """Write the Maturity Assessment sheet"""
        ws = wb.create_sheet("Maturity Assessment")

        org = data.get('organization', {})
        maturity = data.get('maturity', {})
        fields = maturity.get('fields_of_action', {})

        row = 1

        # Title
        ws.merge_cells('A1:D1')
        ws['A1'] = 'SE Maturity Assessment - 4 Fields of Action'
        ws['A1'].font = self.TITLE_FONT
        ws['A1'].alignment = Alignment(horizontal='center')
        row = 3

        # Overall score
        ws[f'A{row}'] = 'Overall Maturity Score:'
        ws[f'A{row}'].font = self.LABEL_FONT
        ws[f'B{row}'] = f"{maturity.get('overall_score', org.get('maturity_score', 0)):.1f}/100"
        row += 1

        ws[f'A{row}'] = 'Maturity Level:'
        ws[f'A{row}'].font = self.LABEL_FONT
        maturity_level = maturity.get('maturity_level', org.get('maturity_level', 1))
        maturity_name = maturity.get('maturity_name', '')
        ws[f'B{row}'] = f"Level {maturity_level}/5" + (f" ({maturity_name})" if maturity_name else "")
        row += 1

        ws[f'A{row}'] = 'Assessment Pathway:'
        ws[f'A{row}'].font = self.LABEL_FONT
        pathway = maturity.get('assessment_pathway', 'Task-based competency assessment')
        ws[f'B{row}'] = pathway
        row += 2

        # Fields of Action table
        ws[f'A{row}'] = '4 Fields of Action Breakdown'
        ws[f'A{row}'].font = self.SECTION_FONT
        row += 1

        headers = ['Field of Action', 'Level', 'Assessment']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = self.HEADER_FONT
            cell.fill = self.HEADER_FILL
            cell.border = self.THIN_BORDER
            cell.alignment = Alignment(horizontal='center')
        row += 1

        for field_name, field_data in fields.items():
            ws.cell(row=row, column=1, value=field_name).border = self.THIN_BORDER

            # Use correct max value (SE Roles has 6 levels, others have 5)
            answer_value = field_data.get('answer_value', 0)
            max_value = field_data.get('max_value', 5)
            score_cell = ws.cell(row=row, column=2, value=f"{answer_value}/{max_value}")
            score_cell.border = self.THIN_BORDER
            score_cell.alignment = Alignment(horizontal='center')

            # Show the description/assessment text
            description = field_data.get('description', '')
            ws.cell(row=row, column=3, value=description).border = self.THIN_BORDER
            row += 1

        # Set column widths
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 50

    def _write_roles_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """Write the Organization Roles sheet"""
        ws = wb.create_sheet("Organization Roles")

        roles = data.get('roles', [])

        row = 1

        # Title
        ws.merge_cells('A1:D1')
        ws['A1'] = 'Organization Roles Inventory'
        ws['A1'].font = self.TITLE_FONT
        ws['A1'].alignment = Alignment(horizontal='center')
        row = 3

        ws[f'A{row}'] = f'Total Roles: {len(roles)}'
        ws[f'A{row}'].font = self.LABEL_FONT
        row += 2

        # Set minimum col A width for overview labels before table '#' column narrows it
        ws.column_dimensions['A'].width = 16

        if not roles:
            ws[f'A{row}'] = 'No roles defined for this organization.'
            return

        # Roles table
        headers = ['#', 'Role Name', 'SE Cluster Mapping', 'Training Program']
        col_widths = [4, 35, 28, 28]

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = self.HEADER_FONT
            cell.fill = self.HEADER_FILL
            cell.border = self.THIN_BORDER
            cell.alignment = Alignment(horizontal='center')
            _cl = get_column_letter(col)
            ws.column_dimensions[_cl].width = max(col_widths[col - 1], ws.column_dimensions[_cl].width or 0)
        row += 1

        for idx, role in enumerate(roles, 1):
            ws.cell(row=row, column=1, value=idx).border = self.THIN_BORDER
            ws.cell(row=row, column=2, value=role.get('name', '')).border = self.THIN_BORDER
            ws.cell(row=row, column=3, value=role.get('se_cluster_name', '-')).border = self.THIN_BORDER
            ws.cell(row=row, column=4, value=role.get('training_program', '-')).border = self.THIN_BORDER
            row += 1

    def _write_modules_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """Write the Training Modules sheet"""
        ws = wb.create_sheet("Training Modules")

        phase3 = data.get('phase3', {})
        modules = phase3.get('modules', [])
        config = phase3.get('config', {})
        view_type = config.get('selected_view', 'competency_level')

        row = 1

        # Title
        ws.merge_cells('A1:H1')
        ws['A1'] = 'Training Modules Overview'
        ws['A1'].font = self.TITLE_FONT
        ws['A1'].alignment = Alignment(horizontal='center')
        row = 3

        ws[f'A{row}'] = 'Training View:'
        ws[f'A{row}'].font = self.LABEL_FONT
        ws[f'B{row}'] = 'Role-Clustered' if view_type == 'role_clustered' else 'Competency-Level'
        row += 1

        ws[f'A{row}'] = 'Total Modules:'
        ws[f'A{row}'].font = self.LABEL_FONT
        ws[f'B{row}'] = len(modules)
        row += 2

        # Set minimum col A width for overview labels before table '#' column narrows it
        ws.column_dimensions['A'].width = 18

        if not modules:
            ws[f'A{row}'] = 'No training modules configured.'
            return

        # Headers based on view type
        if view_type == 'role_clustered':
            headers = ['#', 'Training Program', 'Type', 'Module', 'Level', 'Format', 'Est. Participants', 'Est. Duration (h)']
            col_widths = [4, 22, 12, 35, 14, 28, 14, 14]
        else:
            headers = ['#', 'Module', 'Level', 'Format', 'Est. Participants', 'Est. Duration (h)']
            col_widths = [4, 40, 14, 28, 14, 14]

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = self.HEADER_FONT
            cell.fill = self.HEADER_FILL
            cell.border = self.THIN_BORDER
            cell.alignment = Alignment(horizontal='center')
            _cl = get_column_letter(col)
            ws.column_dimensions[_cl].width = max(col_widths[col - 1], ws.column_dimensions[_cl].width or 0)
        row += 1

        # Style constants for section headers
        COMP_HEADER_FILL = PatternFill(start_color='F2F2F2', end_color='F2F2F2', fill_type='solid')
        COMP_HEADER_FONT = Font(bold=True, italic=True, size=10)
        PKG_ENGINEERS_FILL = PatternFill(start_color='D6E4F0', end_color='D6E4F0', fill_type='solid')
        PKG_MANAGERS_FILL = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')
        PKG_PARTNERS_FILL = PatternFill(start_color='E2EFDA', end_color='E2EFDA', fill_type='solid')
        PKG_HEADER_FONT = Font(bold=True, size=11)
        SUBCLUSTER_FILL = PatternFill(start_color='F5F5F5', end_color='F5F5F5', fill_type='solid')
        SUBCLUSTER_FONT = Font(bold=True, italic=True, size=10, color='606060')

        pkg_fills = {
            'SE for Engineers': PKG_ENGINEERS_FILL,
            'SE for Managers': PKG_MANAGERS_FILL,
            'SE for Interfacing Partners': PKG_PARTNERS_FILL,
        }

        num_cols = len(headers)

        def write_section_hdr(row_num, label, fill, font):
            merge_range = f'A{row_num}:{get_column_letter(num_cols)}{row_num}'
            ws.merge_cells(merge_range)
            cell = ws[f'A{row_num}']
            cell.value = label
            cell.fill = fill
            cell.font = font
            cell.alignment = Alignment(vertical='center')
            return row_num + 1

        def format_module_name(m):
            pmt = m.get('pmt_type', '')
            name = m.get('competency_name', '')
            if pmt and pmt.lower() not in ['combined', '', 'null']:
                name += f" ({pmt.capitalize()})"
            return name

        def resolve_format(m):
            selected_format = m.get('selected_format') or {}
            return (m.get('learning_format_name') or
                    selected_format.get('format_name') or
                    m.get('selected_format_name') or
                    m.get('format_name') or
                    m.get('learning_format') or '-')

        def write_module_row_rc(row_num, m, idx):
            cluster_name = m.get('cluster_name', '-')
            subcluster = m.get('subcluster')
            if 'Engineer' in (cluster_name or ''):
                module_type = 'Common Base' if subcluster == 'common' else 'Role-Specific'
            else:
                module_type = '-'
            level_num = m.get('target_level', 0)
            level_name = self.LEVEL_NAMES.get(level_num, '')
            duration = self.LEVEL_DURATION_HOURS.get(level_num, 2)

            ws.cell(row=row_num, column=1, value=idx).border = self.THIN_BORDER
            ws.cell(row=row_num, column=2, value=cluster_name).border = self.THIN_BORDER
            type_cell = ws.cell(row=row_num, column=3, value=module_type)
            type_cell.border = self.THIN_BORDER
            if module_type == 'Common Base':
                type_cell.fill = self.GREEN_FILL
            elif module_type == 'Role-Specific':
                type_cell.fill = self.LIGHT_FILL
            ws.cell(row=row_num, column=4, value=format_module_name(m)).border = self.THIN_BORDER
            ws.cell(row=row_num, column=5, value=level_name).border = self.THIN_BORDER
            ws.cell(row=row_num, column=6, value=resolve_format(m)).border = self.THIN_BORDER
            ws.cell(row=row_num, column=7, value=m.get('estimated_participants', 0)).border = self.THIN_BORDER
            dur_cell = ws.cell(row=row_num, column=8, value=duration)
            dur_cell.border = self.THIN_BORDER
            dur_cell.alignment = Alignment(horizontal='center')
            return row_num + 1

        def write_module_row_cl(row_num, m, idx):
            level_num = m.get('target_level', 0)
            level_name = self.LEVEL_NAMES.get(level_num, '')
            duration = self.LEVEL_DURATION_HOURS.get(level_num, 2)
            ws.cell(row=row_num, column=1, value=idx).border = self.THIN_BORDER
            ws.cell(row=row_num, column=2, value=format_module_name(m)).border = self.THIN_BORDER
            ws.cell(row=row_num, column=3, value=level_name).border = self.THIN_BORDER
            ws.cell(row=row_num, column=4, value=resolve_format(m)).border = self.THIN_BORDER
            ws.cell(row=row_num, column=5, value=m.get('estimated_participants', 0)).border = self.THIN_BORDER
            dur_cell = ws.cell(row=row_num, column=6, value=duration)
            dur_cell.border = self.THIN_BORDER
            dur_cell.alignment = Alignment(horizontal='center')
            return row_num + 1

        from itertools import groupby

        def write_comp_groups_rfp(row_num, section_modules, idx_counter, write_fn):
            sorted_by_comp = sorted(section_modules, key=lambda m: (m.get('competency_id', m.get('id', 0)), m.get('target_level', 0)))
            for _, grp_iter in groupby(sorted_by_comp, key=lambda m: m.get('competency_id', m.get('id', 0))):
                grp = list(grp_iter)
                comp_name = grp[0].get('competency_name', 'Unknown')
                count = len(grp)
                label = f"    {comp_name} ({count} sub-module{'s' if count != 1 else ''})"
                row_num = write_section_hdr(row_num, label, COMP_HEADER_FILL, COMP_HEADER_FONT)
                for mod in grp:
                    idx_counter[0] += 1
                    row_num = write_fn(row_num, mod, idx_counter[0])
            return row_num

        # Sort modules - unified key: cluster > subcluster > competency_id > target_level
        if view_type == 'role_clustered':
            cluster_order = {'SE for Engineers': 0, 'SE for Managers': 1, 'SE for Interfacing Partners': 2}
            sorted_modules = sorted(modules, key=lambda m: (
                cluster_order.get(m.get('cluster_name', ''), 99),
                m.get('subcluster') or 'z',
                m.get('competency_id', m.get('id', 0)),
                m.get('target_level', 0),
            ))

            idx_counter = [0]
            for cluster_name, cluster_iter in groupby(sorted_modules, key=lambda m: m.get('cluster_name', 'Uncategorized')):
                cluster_mods = list(cluster_iter)
                pkg_fill = pkg_fills.get(cluster_name, PatternFill(start_color='F2F2F2', end_color='F2F2F2', fill_type='solid'))
                row = write_section_hdr(row, cluster_name, pkg_fill, PKG_HEADER_FONT)

                if 'Engineer' in cluster_name:
                    common_mods = [m for m in cluster_mods if m.get('subcluster') == 'common']
                    pathway_mods = [m for m in cluster_mods if m.get('subcluster') != 'common']

                    if common_mods:
                        cb_label = f"  Common Base ({len(common_mods)} module{'s' if len(common_mods) != 1 else ''})"
                        row = write_section_hdr(row, cb_label, SUBCLUSTER_FILL, SUBCLUSTER_FONT)
                        row = write_comp_groups_rfp(row, common_mods, idx_counter, write_module_row_rc)

                    if pathway_mods:
                        rs_label = f"  Role-Specific Pathways ({len(pathway_mods)} module{'s' if len(pathway_mods) != 1 else ''})"
                        row = write_section_hdr(row, rs_label, SUBCLUSTER_FILL, SUBCLUSTER_FONT)
                        row = write_comp_groups_rfp(row, pathway_mods, idx_counter, write_module_row_rc)
                else:
                    row = write_comp_groups_rfp(row, cluster_mods, idx_counter, write_module_row_rc)
        else:
            sorted_modules = sorted(modules, key=lambda m: (m.get('competency_id', m.get('id', 0)), m.get('target_level', 0)))
            idx_counter = [0]
            for _, grp_iter in groupby(sorted_modules, key=lambda m: m.get('competency_id', m.get('id', 0))):
                grp = list(grp_iter)
                comp_name = grp[0].get('competency_name', 'Unknown')
                count = len(grp)
                label = f"{comp_name} ({count} sub-module{'s' if count != 1 else ''})"
                row = write_section_hdr(row, label, COMP_HEADER_FILL, COMP_HEADER_FONT)
                for mod in grp:
                    idx_counter[0] += 1
                    row = write_module_row_cl(row, mod, idx_counter[0])

    def _write_gaps_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """Write the Gap Analysis sheet"""
        ws = wb.create_sheet("Gap Analysis")

        gaps = data.get('gaps', {})

        row = 1

        # Title
        ws.merge_cells('A1:E1')
        ws['A1'] = 'Competency Gap Analysis'
        ws['A1'].font = self.TITLE_FONT
        ws['A1'].alignment = Alignment(horizontal='center')
        row = 3

        # Summary
        ws[f'A{row}'] = 'Total Competency Gaps:'
        ws[f'A{row}'].font = self.LABEL_FONT
        ws[f'B{row}'] = gaps.get('total_gaps', 0)
        row += 2

        # Gaps by Area
        ws[f'A{row}'] = 'Gaps by Competency Area'
        ws[f'A{row}'].font = self.SECTION_FONT
        row += 1

        by_area = gaps.get('by_area', {})
        headers = ['Area', 'Gap Count']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = self.HEADER_FONT
            cell.fill = self.HEADER_FILL
            cell.border = self.THIN_BORDER
        row += 1

        for area, count in by_area.items():
            ws.cell(row=row, column=1, value=area).border = self.THIN_BORDER
            ws.cell(row=row, column=2, value=count).border = self.THIN_BORDER
            row += 1

        row += 1

        # Gaps by Level
        ws[f'A{row}'] = 'Gaps by Target Level'
        ws[f'A{row}'].font = self.SECTION_FONT
        row += 1

        by_level = gaps.get('by_level', {})
        headers = ['Target Level', 'Gap Count']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = self.HEADER_FONT
            cell.fill = self.HEADER_FILL
            cell.border = self.THIN_BORDER
        row += 1

        for level, count in sorted(by_level.items()):
            level_name = self.LEVEL_NAMES.get(int(level), f'Level {level}')
            ws.cell(row=row, column=1, value=f"{level_name} (Level {level})").border = self.THIN_BORDER
            ws.cell(row=row, column=2, value=count).border = self.THIN_BORDER
            row += 1

        row += 1

        # Gaps by Competency
        ws[f'A{row}'] = 'Gaps by Competency'
        ws[f'A{row}'].font = self.SECTION_FONT
        row += 1

        by_comp = gaps.get('by_competency', {})
        if by_comp:
            headers = ['Competency', 'Area', 'Users with Gaps']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = self.HEADER_FONT
                cell.fill = self.HEADER_FILL
                cell.border = self.THIN_BORDER
            row += 1

            for comp_name, comp_data in sorted(by_comp.items()):
                ws.cell(row=row, column=1, value=comp_name).border = self.THIN_BORDER
                ws.cell(row=row, column=2, value=comp_data.get('area', '')).border = self.THIN_BORDER
                ws.cell(row=row, column=3, value=comp_data.get('total_users', 0)).border = self.THIN_BORDER
                row += 1

        # Set column widths
        ws.column_dimensions['A'].width = 35
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 15

    def _write_training_schedule_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """Write the Training Schedule sheet (replaces Timeline sheet).

        Builds pseudo-plan dicts from Phase 3 modules and uses AVIVA service's
        scheduling logic to generate a bin-packed daily schedule with milestones.
        """
        from app.services.phase4_aviva_service import Phase4AvivaService

        phase3 = data.get('phase3', {})
        modules = phase3.get('modules', [])
        config = phase3.get('config', {})
        view_type = config.get('selected_view', 'competency_level')
        timeline = phase3.get('timeline', {})
        milestones = timeline.get('milestones', [])

        # Build pseudo-plan dicts for the schedule generator
        pseudo_plans = []
        for m in modules:
            level = m.get('target_level', 0)
            duration_hours = Phase4AvivaService.LEVEL_DURATION_HOURS.get(level, 2)
            level_name = Phase4AvivaService.LEVEL_NAMES.get(level, f'Level {level}')
            pmt_type = m.get('pmt_type', 'combined')

            # Build module name
            comp_name = m.get('competency_name', '')
            if pmt_type and pmt_type.lower() not in ('combined', '', 'null', 'none'):
                module_name = f"{comp_name} - {level_name} ({pmt_type.capitalize()})"
            else:
                module_name = f"{comp_name} - {level_name}"

            # Resolve format name (same logic as _write_modules_sheet)
            selected_format = m.get('selected_format') or {}
            learning_format = (m.get('learning_format_name') or
                               selected_format.get('format_name') or
                               m.get('selected_format_name') or
                               m.get('format_name') or
                               m.get('learning_format') or '-')

            pseudo_plans.append({
                'module_name': module_name,
                'competency_name': comp_name,
                'competency_id': m.get('competency_id', m.get('id', 0)),
                'target_level': level,
                'total_duration_minutes': duration_hours * 60,
                'learning_format': learning_format,
                'cluster_name': m.get('cluster_name'),
                'subcluster': m.get('subcluster'),
                'pmt_type': pmt_type,
                'estimated_participants': m.get('estimated_participants', 0),
                'roles_needing_training': m.get('roles_needing_training', []),
            })

        # Use AVIVA service's scheduling logic
        aviva_service = Phase4AvivaService(self.db)
        schedule_data = aviva_service._generate_daily_schedule(pseudo_plans, view_type, milestones)

        ws = wb.create_sheet("Training Schedule")
        aviva_service._write_schedule_sheet(ws, schedule_data, view_type, milestones=milestones)

    def _write_role_competency_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """Write the Role-Competency Matrix sheet"""
        ws = wb.create_sheet("Role-Competency Matrix")

        matrix = data.get('role_competency_matrix', [])
        roles = data.get('roles', [])

        row = 1

        # Title
        ws.merge_cells('A1:F1')
        ws['A1'] = 'Role-Competency Matrix (Required Levels)'
        ws['A1'].font = self.TITLE_FONT
        ws['A1'].alignment = Alignment(horizontal='center')
        row = 3

        if not matrix:
            ws[f'A{row}'] = 'No role-competency data available.'
            return

        # Build matrix structure
        role_names = list(set(m['role_name'] for m in matrix))
        comp_names = list(set(m['competency_name'] for m in matrix))

        # Create lookup
        lookup = {}
        for m in matrix:
            key = (m['role_name'], m['competency_name'])
            lookup[key] = m['required_level']

        # Headers
        ws.cell(row=row, column=1, value='Competency').font = self.HEADER_FONT
        ws.cell(row=row, column=1).fill = self.HEADER_FILL
        ws.cell(row=row, column=1).border = self.THIN_BORDER

        for col, role_name in enumerate(role_names, 2):
            cell = ws.cell(row=row, column=col, value=role_name)
            cell.font = self.HEADER_FONT
            cell.fill = self.HEADER_FILL
            cell.border = self.THIN_BORDER
            cell.alignment = Alignment(horizontal='center', wrap_text=True)
        row += 1

        # Data rows
        for comp_name in sorted(comp_names):
            ws.cell(row=row, column=1, value=comp_name).border = self.THIN_BORDER

            for col, role_name in enumerate(role_names, 2):
                level = lookup.get((role_name, comp_name), 0)
                cell = ws.cell(row=row, column=col, value=level if level > 0 else '-')
                cell.border = self.THIN_BORDER
                cell.alignment = Alignment(horizontal='center')

                # Color coding
                if level >= 4:
                    cell.fill = self.GREEN_FILL
                elif level >= 2:
                    cell.fill = self.YELLOW_FILL
            row += 1

        # Set column widths
        ws.column_dimensions['A'].width = 30
        for col in range(2, len(role_names) + 2):
            ws.column_dimensions[get_column_letter(col)].width = 20

    def _write_aviva_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """Write the AVIVA Plans sheet"""
        ws = wb.create_sheet("AVIVA Plans")

        plans = data.get('aviva_plans', [])
        phase3 = data.get('phase3', {})
        config = phase3.get('config', {})
        view_type = config.get('selected_view', 'competency_level')

        row = 1

        # Title
        ws.merge_cells('A1:G1')
        ws['A1'] = 'AVIVA Didactic Plans'
        ws['A1'].font = self.TITLE_FONT
        ws['A1'].alignment = Alignment(horizontal='center')
        row = 3

        if not plans:
            ws[f'A{row}'] = 'No AVIVA plans generated.'
            return

        # Sort plans
        if view_type == 'role_clustered':
            cluster_order = {'SE for Engineers': 0, 'SE for Managers': 1, 'SE for Interfacing Partners': 2}
            sorted_plans = sorted(plans, key=lambda p: (
                cluster_order.get(p.get('cluster_name', ''), 99),
                p.get('subcluster') or 'z',
                p.get('competency_id', 0)
            ))
        else:
            sorted_plans = sorted(plans, key=lambda p: p.get('competency_id', 0))

        # AVIVA headers
        aviva_headers = ['Start', 'Min', 'Type', 'AVIVA', 'What (Content)', 'How (Method)', 'Material']
        aviva_col_widths = [18, 8, 10, 15, 50, 25, 45]

        for col, width in enumerate(aviva_col_widths, 1):
            ws.column_dimensions[get_column_letter(col)].width = width

        # Write each module
        for plan in sorted_plans:
            level_name = self.LEVEL_NAMES.get(plan.get('target_level', 0), '')

            # Module header
            ws.merge_cells(f'A{row}:G{row}')
            cell = ws[f'A{row}']
            cell.value = f"Module: {plan.get('module_name', 'Unknown')}"
            cell.font = Font(bold=True, size=11, color="FFFFFF")
            cell.fill = self.SUBHEADER_FILL
            cell.alignment = Alignment(horizontal='left', vertical='center')
            row += 1

            # Module details
            if view_type == 'role_clustered' and plan.get('cluster_name'):
                ws[f'A{row}'] = 'Training Program:'
                ws[f'A{row}'].font = self.LABEL_FONT
                ws[f'B{row}'] = plan.get('cluster_name', '')
                row += 1

            ws[f'A{row}'] = 'Level:'
            ws[f'A{row}'].font = self.LABEL_FONT
            ws[f'B{row}'] = level_name
            ws[f'C{row}'] = 'Format:'
            ws[f'C{row}'].font = self.LABEL_FONT
            ws[f'D{row}'] = plan.get('learning_format', 'N/A')
            ws[f'E{row}'] = 'Duration:'
            ws[f'E{row}'].font = self.LABEL_FONT
            ws[f'F{row}'] = f"{plan.get('total_duration_minutes', 0)} min"
            row += 1

            # Learning objective
            ws[f'A{row}'] = 'Learning Objective:'
            ws[f'A{row}'].font = self.LABEL_FONT
            ws.merge_cells(f'B{row}:G{row}')
            lo_cell = ws[f'B{row}']
            lo_cell.value = plan.get('learning_objective', 'N/A')
            lo_cell.alignment = Alignment(wrap_text=True)
            row += 1

            # Content topics
            topics = plan.get('content_topics', [])
            if topics:
                ws[f'A{row}'] = 'Content Topics:'
                ws[f'A{row}'].font = self.LABEL_FONT
                ws.merge_cells(f'B{row}:G{row}')
                ws[f'B{row}'] = ', '.join(topics)
                ws[f'B{row}'].alignment = Alignment(wrap_text=True)
                row += 1

            row += 1  # Empty row

            # AVIVA table headers
            for col, header in enumerate(aviva_headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = self.HEADER_FONT
                cell.fill = self.HEADER_FILL
                cell.border = self.THIN_BORDER
                cell.alignment = Alignment(horizontal='center')
            row += 1

            # AVIVA activities
            activities = plan.get('activities', [])
            for activity in activities:
                ws.cell(row=row, column=1, value=activity.get('start_time', '')).border = self.THIN_BORDER
                ws.cell(row=row, column=2, value=activity.get('duration_min', 0)).border = self.THIN_BORDER
                ws.cell(row=row, column=3, value=activity.get('type', '')).border = self.THIN_BORDER
                ws.cell(row=row, column=4, value=activity.get('aviva_phase', '')).border = self.THIN_BORDER

                content_cell = ws.cell(row=row, column=5, value=activity.get('content', ''))
                content_cell.border = self.THIN_BORDER
                content_cell.alignment = Alignment(wrap_text=True)

                ws.cell(row=row, column=6, value=activity.get('method', '')).border = self.THIN_BORDER
                ws.cell(row=row, column=7, value=activity.get('material', '')).border = self.THIN_BORDER

                for col in [1, 2, 3, 4]:
                    ws.cell(row=row, column=col).alignment = Alignment(horizontal='center')
                row += 1

            row += 3  # Space between modules

    # =========================================================================
    # WORD DOCUMENT EXPORT (LLM-ENHANCED)
    # =========================================================================

    def export_rfp_to_word(self, organization_id: int, include_llm: bool = True,
                           module_ids: Optional[List[int]] = None) -> io.BytesIO:
        """
        Export comprehensive RFP document to Word format with LLM-generated content.

        Creates a professional document following the ZEISS RFP reference structure:
        - Part 1: Context and Core Concept
        - Part 2: Service Requirements
        - Part 3: Training Program
        - Part 4: Training Module Details

        Uses parallel LLM generation to speed up module content creation.

        Args:
            organization_id: Organization ID
            include_llm: Whether to use LLM for content generation
            module_ids: Optional list of module IDs to filter (from AVIVA selection)
        """
        current_app.logger.info(f"[RFP] Starting Word export for organization {organization_id}, include_llm={include_llm}")

        # Get all data (with optional module filtering)
        data = self.get_rfp_data(organization_id, include_aviva=False, module_ids=module_ids)
        org = data.get('organization', {})
        phase3 = data.get('phase3', {})
        modules = phase3.get('modules', [])
        pmt_context = data.get('pmt_context', {})

        # Pre-generate all module content in PARALLEL if LLM is enabled
        # This significantly speeds up the export (from sequential to parallel LLM calls)
        pre_generated_content = None
        if include_llm and modules:
            current_app.logger.info(f"[RFP] Starting parallel LLM generation for {len(modules)} modules")
            # Sort modules first to match the order in document
            sorted_modules = sorted(modules, key=lambda m: (
                m.get('competency_id', 0),
                m.get('target_level', 0)
            ))
            pre_generated_content = self.generate_all_module_content_parallel(
                sorted_modules,
                pmt_context,
                max_workers=5  # Limit to avoid API rate limits
            )
            current_app.logger.info(f"[RFP] Parallel generation complete")

        # Create Word document
        doc = Document()

        # Set up document styles
        self._setup_word_styles(doc)

        # Add document sections
        self._add_title_page(doc, data)

        # GenAI Disclaimer
        disclaimer_para = doc.add_paragraph()
        disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        disclaimer_run = disclaimer_para.add_run(
            "Disclaimer: This document contains content generated using Generative AI (GenAI) "
            "and may contain inaccuracies or errors. It is intended as a reference document and "
            "starting point. Review and modifications may be necessary to ensure accuracy and "
            "alignment with your specific organizational requirements."
        )
        disclaimer_run.italic = True
        disclaimer_run.font.size = Pt(9)
        disclaimer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        disclaimer_para.space_after = Pt(12)

        self._add_table_of_contents(doc)
        self._add_context_section(doc, data, include_llm)
        self._add_service_requirements_section(doc, data)
        self._add_training_program_section(doc, data)
        self._add_module_details_section(doc, data, include_llm, pre_generated_content)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        current_app.logger.info(f"[RFP] Word export completed")
        return buffer

    def _setup_word_styles(self, doc: Document):
        """Set up document styles for consistent formatting"""
        # Get or create styles
        styles = doc.styles

        # Define color constants
        self.HEADER_RGB = RGBColor(0x45, 0x5A, 0x64)  # #455A64
        self.SUBHEADER_RGB = RGBColor(0x60, 0x7D, 0x8B)  # #607D8B

        # Modify default styles
        style = styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Heading 1
        if 'Heading 1' in styles:
            h1 = styles['Heading 1']
            h1.font.name = 'Calibri'
            h1.font.size = Pt(18)
            h1.font.bold = True
            h1.font.color.rgb = self.HEADER_RGB

        # Heading 2
        if 'Heading 2' in styles:
            h2 = styles['Heading 2']
            h2.font.name = 'Calibri'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.font.color.rgb = self.SUBHEADER_RGB

        # Heading 3
        if 'Heading 3' in styles:
            h3 = styles['Heading 3']
            h3.font.name = 'Calibri'
            h3.font.size = Pt(12)
            h3.font.bold = True

    def _add_phase_origin(self, paragraph, text: str):
        """Add a phase origin annotation (italic, gray) to a paragraph"""
        run = paragraph.add_run(f"  [{text}]")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    def _add_intro_paragraph(self, doc: Document, text: str):
        """Add an introductory/descriptive paragraph in slightly smaller, gray text"""
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        p.space_after = Pt(6)
        return p

    def _add_title_page(self, doc: Document, data: Dict[str, Any]):
        """Add title page to the document"""
        org = data.get('organization', {})

        # Add some spacing at top
        for _ in range(3):
            doc.add_paragraph()

        # Main title
        title = doc.add_paragraph()
        title_run = title.add_run("Systems Engineering Qualification Program")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = self.HEADER_RGB
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run("Input for identifying training service providers")
        subtitle_run.font.size = Pt(18)
        subtitle_run.font.color.rgb = self.SUBHEADER_RGB
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing
        for _ in range(4):
            doc.add_paragraph()

        # Organization name
        org_para = doc.add_paragraph()
        org_run = org_para.add_run(org.get('name', 'Organization'))
        org_run.font.size = Pt(16)
        org_run.font.bold = True
        org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(datetime.now().strftime('%d %B %Y'))
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page break
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document):
        """Add table of contents placeholder"""
        doc.add_heading('Contents', level=1)

        # Add TOC items (simplified - Word will auto-generate if needed)
        toc_items = [
            "Context and Core Concept",
            "    Systems Engineering Context",
            "    Core Concept",
            "    Qualification Strategy",
            "Service Requirements",
            "    Service Description - General",
            "    Requirements/Constraints",
            "    Personnel Subject to Training",
            "Training Program",
            "    Key Competencies",
            "    Training Structure Overview",
            "    Timeline",
            "Training Module Details"
        ]

        for item in toc_items:
            p = doc.add_paragraph(item)
            p.paragraph_format.left_indent = Inches(0.5) if item.startswith("    ") else Inches(0)

        doc.add_page_break()

    def _add_context_section(self, doc: Document, data: Dict[str, Any], include_llm: bool):
        """Add Part 1: Context and Core Concept section"""
        doc.add_heading('Part 1: Context and Core Concept', level=1)

        # 1.1 Systems Engineering Context
        doc.add_heading('1.1 Systems Engineering Context', level=2)

        context_text = """Systems Engineering requires all the capabilities needed to successfully work on development projects within the system level activities of the V-Model, as well as for concept and advanced development.

This includes requirements engineering, architectural modelling, test specification, risk management, design considerations (value, reliability, sustainability, security), and integration with component-level engineering disciplines (electrical, mechanical, software)."""

        doc.add_paragraph(context_text)

        # 1.2 Core Concept
        doc.add_heading('1.2 Core Concept', level=2)

        if include_llm:
            # Generate Core Concept using LLM
            core_concept = self.generate_core_concept(data)
        else:
            core_concept = self._get_core_concept_fallback(data)

        # Add the core concept text
        for paragraph in core_concept.split('\n\n'):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph.strip())

        # 1.3 Qualification Strategy
        doc.add_heading('1.3 Qualification Strategy', level=2)

        strategies = data.get('strategies', [])
        if strategies:
            self._add_intro_paragraph(doc,
                "The following qualification strategies were selected during the strategic planning phase "
                "to guide the design of the training program. Each strategy defines the approach, "
                "target audience, and qualification depth for different segments of the organization.")

            from app.strategy_selection_engine import SE_TRAINING_STRATEGIES
            strategy_detail_map = {s['name']: s for s in SE_TRAINING_STRATEGIES}

            for s in strategies:
                s_name = s.get('name', '')
                is_primary = s.get('is_primary', False)

                # Strategy name heading
                label = f"{s_name} (Primary Strategy)" if is_primary else s_name
                p = doc.add_paragraph()
                run = p.add_run(label)
                run.bold = True
                run.font.size = Pt(11)

                # Description
                desc = s.get('description', '') or strategy_detail_map.get(s_name, {}).get('description', '')
                if desc:
                    doc.add_paragraph(desc)

                # Key details
                detail = strategy_detail_map.get(s_name, {})
                details_parts = []
                if detail.get('targetAudience'):
                    details_parts.append(f"Target Audience: {detail['targetAudience']}")
                if detail.get('qualificationLevel'):
                    details_parts.append(f"Qualification Level: {detail['qualificationLevel']}")
                if detail.get('suitablePhase'):
                    details_parts.append(f"Suitable Phase: {detail['suitablePhase']}")
                if detail.get('duration'):
                    details_parts.append(f"Estimated Duration: {detail['duration']}")
                if details_parts:
                    for dp in details_parts:
                        bullet_p = doc.add_paragraph(dp, style='List Bullet')
                        for r in bullet_p.runs:
                            r.font.size = Pt(10)

                doc.add_paragraph()  # Spacing between strategies
        else:
            doc.add_paragraph("No qualification strategy has been selected.")

        doc.add_page_break()

    def _add_service_requirements_section(self, doc: Document, data: Dict[str, Any]):
        """Add Part 2: Service Requirements section"""
        doc.add_heading('Part 2: Service Requirements', level=1)

        # 2.1 Service Description - General
        doc.add_heading('2.1 Service Description - General', level=2)

        self._add_intro_paragraph(doc,
            "The following describes the general scope of services expected from the training provider. "
            "These requirements apply to all training modules defined in this RFP.")

        general_items = [
            "Concepts for training modules for pre-defined contents; structure/modality didactically optimized.",
            "Creation of training material (English).",
            "Alignment with the organization regarding training contents and modalities; review of material.",
            "Provision of training; the trainings shall include hands-on exercises with the given tools.",
            "Reviewing the correct and goal-driven application of the training participants for certification."
        ]

        for item in general_items:
            p = doc.add_paragraph(item, style='List Bullet')

        # 2.2 Requirements/Constraints
        doc.add_heading('2.2 Requirements/Constraints', level=2)

        self._add_intro_paragraph(doc,
            "The following requirements and constraints were derived from the organization profile, "
            "selected learning formats, and timeline planning defined in Phases 1-3.")

        requirements = self.generate_service_requirements(data)
        for req in requirements:
            p = doc.add_paragraph(req, style='List Bullet')

        # 2.3 Personnel Subject to Training
        doc.add_heading('2.3 Personnel Subject to Training', level=2)

        roles = data.get('roles', [])
        target_group = data.get('target_group', {})

        self._add_intro_paragraph(doc,
            "The target group size was defined during the organization profile setup (Phase 1), "
            "and the roles below were identified during the competency assessment (Phase 2). "
            "Each role has been assigned to a training program cluster based on its SE competency profile.")

        intro = f"Total target group: {target_group.get('size', 'N/A')} employees across {len(roles)} identified roles."
        doc.add_paragraph(intro)

        if roles:
            # Create roles table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            hdr_cells = table.rows[0].cells
            headers = ['Role Name', 'SE Cluster', 'Training Program']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                self._set_cell_shading(hdr_cells[i], "455A64")

            # Data rows - show all roles
            for role in roles:
                row_cells = table.add_row().cells
                row_cells[0].text = role.get('name', '-')
                row_cells[1].text = role.get('se_cluster_name', '-') or '-'
                row_cells[2].text = role.get('training_program', '-') or '-'

            # Set column widths
            for row in table.rows:
                row.cells[0].width = Inches(2.5)
                row.cells[1].width = Inches(2.0)
                row.cells[2].width = Inches(2.0)

        doc.add_page_break()

    def _add_training_program_section(self, doc: Document, data: Dict[str, Any]):
        """Add Part 3: Training Program section"""
        doc.add_heading('Part 3: Training Program', level=1)

        phase3 = data.get('phase3', {})
        modules = phase3.get('modules', [])
        summary = phase3.get('summary', {})
        timeline = phase3.get('timeline', {})
        config = phase3.get('config', {})
        view_type = config.get('selected_view', 'competency_level')
        is_role_clustered = view_type == 'role_clustered'

        # 3.1 Key Competencies
        doc.add_heading('3.1 Key Competencies', level=2)

        self._add_intro_paragraph(doc,
            "The following SE competencies were identified as requiring qualification based on the "
            "competency gap analysis conducted in Phase 2. For each competency, the gap between "
            "current and required competency levels was assessed across all roles, and training "
            "modules were generated accordingly.")

        # Get unique competencies
        competencies = set()
        for m in modules:
            comp_name = m.get('competency_name')
            if comp_name:
                competencies.add(comp_name)

        if competencies:
            p = doc.add_paragraph()
            run = p.add_run(f"{len(competencies)} key competencies")
            run.bold = True
            p.add_run(" are addressed in this qualification program:")
            for comp in sorted(competencies):
                doc.add_paragraph(comp, style='List Bullet')
        else:
            doc.add_paragraph("No competencies defined.")

        # 3.2 Training Structure Overview
        doc.add_heading('3.2 Training Structure Overview', level=2)

        if is_role_clustered:
            view_desc = (
                "The training program uses a Role-Clustered structure, grouping modules into "
                "three training packages based on role assignments: SE for Engineers (with Common Base "
                "and Role-Specific Pathways), SE for Managers, and SE for Interfacing Partners. "
                "This structure was selected in Phase 3 based on the organization's maturity level and role diversity."
            )
        else:
            view_desc = (
                "The training program uses a Competency-Level structure, grouping modules by the "
                "16 SE competency areas. Each competency group contains sub-modules at identified "
                "qualification levels (Knowing, Understanding, Applying, Mastering). "
                "This structure was selected in Phase 3."
            )
        self._add_intro_paragraph(doc, view_desc)

        # Summary statistics
        total_duration = sum(self.LEVEL_DURATION_HOURS.get(m.get('target_level', 0), 2) for m in modules)
        duration_display = f"{total_duration} hours"
        if total_duration >= 8:
            duration_display += f" (~{total_duration / 8:.1f} training days)"

        p = doc.add_paragraph()
        p.add_run("Total Training Modules: ").bold = True
        p.add_run(str(len(modules)))

        p = doc.add_paragraph()
        p.add_run("Total Estimated Duration: ").bold = True
        p.add_run(duration_display)

        p = doc.add_paragraph()
        p.add_run("Training View: ").bold = True
        p.add_run('Role-Clustered' if is_role_clustered else 'Competency-Level')

        format_dist = summary.get('format_distribution', {})
        if format_dist:
            p = doc.add_paragraph()
            p.add_run("Format Distribution: ").bold = True
            p.add_run(', '.join(f"{fmt}: {count}" for fmt, count in format_dist.items()))

        # Modules table
        if modules:
            doc.add_paragraph()  # Spacing

            # Sort modules based on view type
            if is_role_clustered:
                cluster_order = {'SE for Engineers': 0, 'SE for Managers': 1, 'SE for Interfacing Partners': 2}
                sorted_modules = sorted(modules, key=lambda m: (
                    cluster_order.get(m.get('cluster_name', ''), 99),
                    m.get('subcluster') or 'z',  # 'common' before 'pathway'
                    m.get('competency_id', 0),
                    m.get('target_level', 0)
                ))
            else:
                sorted_modules = sorted(modules, key=lambda m: (
                    m.get('competency_id', 0),
                    m.get('target_level', 0)
                ))

            # Different table structure based on view type
            if is_role_clustered:
                table = doc.add_table(rows=1, cols=8)
                headers = ['Training Program', 'Type', 'Module', 'Level', 'Format', 'Roles', 'Participants', 'Duration (h)']
            else:
                table = doc.add_table(rows=1, cols=6)
                headers = ['Module', 'Level', 'Format', 'Roles', 'Participants', 'Duration (h)']

            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                self._set_cell_shading(hdr_cells[i], "455A64")

            # Data rows - show all modules
            for m in sorted_modules:
                row_cells = table.add_row().cells
                level_name = self.LEVEL_NAMES.get(m.get('target_level', 0), 'N/A')

                # Get format name
                selected_format = m.get('selected_format') or {}
                format_name = (m.get('learning_format_name') or
                              selected_format.get('format_name') or '-')

                # Build module display name with PMT type if applicable
                competency_name = m.get('competency_name', '-')
                pmt_type = m.get('pmt_type', '')
                if pmt_type and pmt_type.lower() not in ['combined', '', 'null', 'none']:
                    module_display_name = f"{competency_name} ({pmt_type.capitalize()})"
                else:
                    module_display_name = competency_name

                # Get roles
                roles = m.get('roles_needing_training', [])
                roles_str = ', '.join(roles) if roles and isinstance(roles, list) else '-'

                duration = str(self.LEVEL_DURATION_HOURS.get(m.get('target_level', 0), 2))

                if is_role_clustered:
                    # Determine module type
                    cluster_name = m.get('cluster_name', '-')
                    subcluster = m.get('subcluster')
                    if 'Engineer' in (cluster_name or ''):
                        module_type = 'Common Base' if subcluster == 'common' else 'Role-Specific'
                    else:
                        module_type = '-'

                    row_cells[0].text = cluster_name
                    row_cells[1].text = module_type
                    row_cells[2].text = module_display_name
                    row_cells[3].text = level_name
                    row_cells[4].text = format_name
                    row_cells[5].text = roles_str
                    row_cells[6].text = str(m.get('estimated_participants', 0))
                    row_cells[7].text = duration
                else:
                    row_cells[0].text = module_display_name
                    row_cells[1].text = level_name
                    row_cells[2].text = format_name
                    row_cells[3].text = roles_str
                    row_cells[4].text = str(m.get('estimated_participants', 0))
                    row_cells[5].text = duration

        # 3.3 Timeline
        doc.add_heading('3.3 Timeline', level=2)

        self._add_intro_paragraph(doc,
            "The implementation timeline below was estimated in Phase 3 based on the organization's "
            "maturity level, target group size, number of training modules, and selected learning formats. "
            "It outlines the key milestones from concept development through full rollout.")

        milestones = timeline.get('milestones', [])
        if milestones:
            for ms in milestones:
                name = ms.get('milestone_name', ms.get('name', ''))
                date = ms.get('estimated_date', '')
                desc = ms.get('milestone_description', ms.get('description', ''))

                p = doc.add_paragraph()
                p.add_run(f"{name}").bold = True
                if date:
                    p.add_run(f" ({date})")
                if desc:
                    p.add_run(f": {desc}")
        else:
            doc.add_paragraph("Timeline to be determined.")

        doc.add_page_break()

    def _add_module_details_section(
        self,
        doc: Document,
        data: Dict[str, Any],
        include_llm: bool,
        pre_generated_content: Optional[Dict[int, Dict[str, Any]]] = None
    ):
        """
        Add Part 4: Training Module Details section with LLM-generated Goals and Contents.

        Args:
            doc: Word document to add content to
            data: RFP data dictionary
            include_llm: Whether to use LLM for generation
            pre_generated_content: Optional dict of pre-generated content from parallel processing
                                   Format: {module_index: {'goals': [...], 'contents': [...]}}
        """
        doc.add_heading('Part 4: Training Module Details', level=1)

        self._add_intro_paragraph(doc,
            "This section provides detailed training module specifications including learning goals "
            "and content outlines. Each module targets a specific SE competency at an identified "
            "qualification level, with goals and contents generated to guide training material development. "
            "The modules are derived from the competency gap analysis (Phase 2) and training structure "
            "planning (Phase 3).")

        phase3 = data.get('phase3', {})
        modules = phase3.get('modules', [])
        pmt = data.get('pmt_context', {})

        if not modules:
            doc.add_paragraph("No training modules configured.")
            return

        # Sort modules (same order as parallel generation)
        sorted_modules = sorted(modules, key=lambda m: (
            m.get('competency_id', 0),
            m.get('target_level', 0)
        ))

        current_app.logger.info(f"[RFP] Adding details for {len(sorted_modules)} modules")

        for idx, module in enumerate(sorted_modules):
            competency_id = module.get('competency_id')
            competency_name = module.get('competency_name', 'Unknown')
            target_level = module.get('target_level', 2)
            level_name = self.LEVEL_NAMES.get(target_level, 'Understanding')

            # Get content topics from baseline (for fallback if needed)
            content_topics = self.get_content_topics_for_competency(competency_id) if competency_id else []

            # Get format name
            selected_format = module.get('selected_format') or {}
            format_name = (module.get('learning_format_name') or
                          selected_format.get('format_name') or 'N/A')

            # Module header (1-indexed for display)
            doc.add_heading(f'{idx + 1}. {competency_name} - Level {level_name}', level=2)

            # Module info paragraph
            duration_hours = self.LEVEL_DURATION_HOURS.get(target_level, 2)
            info = f"Target Level: {level_name} | Format: {format_name} | Est. Duration: {duration_hours}h"
            cluster_name = module.get('cluster_name')
            if cluster_name:
                info += f" | Training Program: {cluster_name}"
            participants = module.get('estimated_participants', 0)
            if participants:
                info += f" | Est. Participants: {participants}"

            info_para = doc.add_paragraph(info)
            info_para.runs[0].font.italic = True
            info_para.runs[0].font.color.rgb = self.SUBHEADER_RGB

            # Roles needing training
            module_roles = module.get('roles_needing_training', [])
            if module_roles and isinstance(module_roles, list):
                roles_para = doc.add_paragraph()
                roles_run = roles_para.add_run(f"Roles: {', '.join(module_roles)}")
                roles_run.font.italic = True
                roles_run.font.size = Pt(10)
                roles_run.font.color.rgb = RGBColor(0x90, 0x93, 0x99)

            # Check if we have pre-generated content for this module
            if pre_generated_content and idx in pre_generated_content:
                # Use pre-generated content (from parallel processing)
                goals = pre_generated_content[idx].get('goals', [])
                contents = pre_generated_content[idx].get('contents', [])
            elif include_llm:
                # Fall back to sequential generation (should rarely happen)
                module_data = {
                    'competency_name': competency_name,
                    'target_level': target_level,
                    'learning_objective': module.get('learning_objective', ''),
                    'content_topics': content_topics,
                    'pmt_type': module.get('pmt_type', 'combined'),
                    'duration_hours': module.get('estimated_duration_hours', 8),
                    'org_tools': pmt.get('tools', ''),
                    'learning_format': format_name
                }
                goals = self.generate_module_goals(module_data)
                contents = self.generate_module_contents(module_data)
            else:
                # Use fallback content (no LLM)
                module_data = {
                    'competency_name': competency_name,
                    'target_level': target_level,
                    'learning_objective': module.get('learning_objective', ''),
                    'content_topics': content_topics,
                    'pmt_type': module.get('pmt_type', 'combined'),
                    'duration_hours': module.get('estimated_duration_hours', 8),
                    'org_tools': pmt.get('tools', ''),
                    'learning_format': format_name
                }
                goals = self._get_module_goals_fallback(module_data)
                contents = self._get_module_contents_fallback(module_data)

            # Goals section
            doc.add_heading('Goals', level=3)
            for goal in goals:
                doc.add_paragraph(goal, style='List Bullet')

            # Contents section
            doc.add_heading('Contents', level=3)
            for content in contents:
                # Check if it's a sub-item (starts with spaces or letters like "a." "b.")
                if content.strip().startswith(('a.', 'b.', 'c.', 'd.', 'e.')) or content.startswith('   '):
                    p = doc.add_paragraph(content.strip())
                    p.paragraph_format.left_indent = Inches(0.5)
                else:
                    doc.add_paragraph(content)

            # Add some spacing between modules
            doc.add_paragraph()

    def _set_cell_shading(self, cell, hex_color: str):
        """Set cell background shading"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), hex_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

        # Set text color to white for dark backgrounds
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # =========================================================================
    # SAVE EXPORT RECORD
    # =========================================================================

    def save_export_record(self, organization_id: int, export_format: str,
                          include_aviva: bool, data_snapshot: Dict = None) -> int:
        """Save export record to database"""
        result = self.db.execute(
            text("""
                INSERT INTO phase4_rfp_export
                    (organization_id, export_format, include_aviva, export_data)
                VALUES (:org_id, :format, :aviva, :data)
                RETURNING id
            """),
            {
                'org_id': organization_id,
                'format': export_format,
                'aviva': include_aviva,
                'data': json.dumps(data_snapshot) if data_snapshot else None
            }
        )
        self.db.commit()
        return result.fetchone().id

    def get_export_history(self, organization_id: int) -> List[Dict[str, Any]]:
        """Get export history for an organization"""
        results = self.db.execute(
            text("""
                SELECT id, export_format, include_aviva, created_at
                FROM phase4_rfp_export
                WHERE organization_id = :org_id
                ORDER BY created_at DESC
                LIMIT 20
            """),
            {'org_id': organization_id}
        ).fetchall()

        return [
            {
                'id': r.id,
                'format': r.export_format,
                'include_aviva': r.include_aviva,
                'created_at': r.created_at.strftime('%Y-%m-%d %H:%M') if r.created_at else None
            }
            for r in results
        ]
