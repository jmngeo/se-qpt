"""
Phase 2 Learning Objectives Routes Blueprint
Handles all Phase 2 Task 3 learning objectives generation and management
"""

from flask import Blueprint, request, jsonify, current_app, send_file, make_response
from flask_jwt_extended import jwt_required, get_jwt_identity
from datetime import datetime
import json
import traceback
import io
import os
import re

from models import (
    db,
    Organization,
    LearningStrategy,
    OrganizationPMTContext,
    GeneratedLearningObjectives,
    User,
    UserAssessment,
    OrganizationExistingTraining,
    Competency
)

# Import Phase 2 Task 3 setup function
import sys
setup_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'setup', 'utils'))
if setup_path not in sys.path:
    sys.path.insert(0, setup_path)
from setup_phase2_task3_for_org import setup_phase2_task3_strategies

# Create blueprint
phase2_learning_bp = Blueprint('phase2_learning', __name__)


# ==============================================================================
# PATH HELPERS (Docker-safe, works from any working directory)
# ==============================================================================
from pathlib import Path

def _get_backend_root():
    """Get backend root path: src/backend/"""
    # This file: src/backend/app/routes/phase2_learning.py
    # Navigate: routes -> app -> backend
    current_file = Path(__file__).resolve()
    return current_file.parent.parent.parent


def _get_project_root():
    """Get project root path using absolute path resolution (Docker-safe)"""
    # Navigate: backend -> src -> project_root
    return _get_backend_root().parent.parent


def _get_pmt_examples_dir():
    """Get PMT reference examples directory path."""
    return str(_get_backend_root() / 'data' / 'templates' / 'pmt_examples')


# ==============================================================================
# GENERATION AND RETRIEVAL ENDPOINTS
# ==============================================================================

@phase2_learning_bp.route('/phase2/learning-objectives/generate', methods=['POST'])
def api_generate_learning_objectives():
    """
    Generate learning objectives for an organization (Week 2 Implementation)

    This endpoint uses the new Design v5 implementation with all 8 algorithms:
    1. Calculate combined targets (separate TTT)
    2. Validate mastery requirements
    3. Detect gaps (role-based or organizational)
    4. Determine training methods
    5. Process TTT gaps
    6. Generate learning objectives (with PMT customization)
    7. Structure pyramid output
    8. Generate strategy comparison

    Request Body:
        {
            "organization_id": 28,
            "selected_strategies": [
                {"strategy_id": 1, "strategy_name": "Continuous support"},
                {"strategy_id": 6, "strategy_name": "Train the trainer"}
            ],
            "pmt_context": {  // Optional
                "processes": "ISO 26262, ASPICE",
                "methods": "Scrum, V-Model",
                "tools": "DOORS, JIRA"
            }
        }

    Response (Success):
        {
            "success": true,
            "data": {
                "main_pyramid": {
                    "levels": {1: {...}, 2: {...}, 4: {...}, 6: {...}},
                    "metadata": {...}
                },
                "train_the_trainer": {...} or null,
                "validation": {...},
                "strategy_comparison": {...}
            },
            "metadata": {
                "organization_id": 28,
                "selected_strategies": [...],
                "pmt_customization": true/false,
                "has_roles": true/false,
                "generation_timestamp": "2025-11-25T...",
                "processing_time_seconds": 0.45
            }
        }

    Response (Error):
        {
            "success": false,
            "error": "Error description",
            "error_type": "INVALID_REQUEST" | "ORGANIZATION_NOT_FOUND" | "INTERNAL_ERROR",
            "details": {...}
        }
    """
    try:
        from app.services.learning_objectives_core import generate_complete_learning_objectives

        data = request.get_json()

        # Validate request
        if not data:
            return jsonify({
                'success': False,
                'error': 'Request body is required',
                'error_type': 'INVALID_REQUEST'
            }), 400

        if 'organization_id' not in data:
            return jsonify({
                'success': False,
                'error': 'Missing organization_id in request body',
                'error_type': 'INVALID_REQUEST'
            }), 400

        if 'selected_strategies' not in data or not isinstance(data['selected_strategies'], list):
            return jsonify({
                'success': False,
                'error': 'Missing or invalid selected_strategies in request body (must be array)',
                'error_type': 'INVALID_REQUEST'
            }), 400

        organization_id = data['organization_id']
        pmt_context = data.get('pmt_context', None)  # Optional - can be from request or DB

        # Validate organization exists
        org = Organization.query.get(organization_id)
        if not org:
            return jsonify({
                'success': False,
                'error': f'Organization {organization_id} not found',
                'error_type': 'ORGANIZATION_NOT_FOUND'
            }), 404

        # IMPORTANT: Always read strategies from DB for consistent hashing
        # The frontend sends strategies, but we re-read from DB to ensure consistent formatting
        selected_strategies_db = LearningStrategy.query.filter_by(
            organization_id=organization_id,
            selected=True
        ).order_by(LearningStrategy.priority.asc()).all()

        if not selected_strategies_db:
            return jsonify({
                'success': False,
                'error': 'No learning strategies selected for this organization',
                'error_type': 'NO_STRATEGIES'
            }), 400

        # Format strategies consistently with GET endpoint
        selected_strategies = [
            {
                'strategy_id': s.strategy_template_id or s.id,
                'strategy_name': s.strategy_name
            }
            for s in selected_strategies_db
        ]

        # If PMT context not provided in request, try to get from database
        if not pmt_context:
            pmt_record = OrganizationPMTContext.query.filter_by(organization_id=organization_id).first()
            if pmt_record and pmt_record.is_complete():
                pmt_context = {
                    'processes': pmt_record.processes,
                    'methods': pmt_record.methods,
                    'tools': pmt_record.tools
                }
                print(f"[api_generate_learning_objectives] PMT loaded from DB: {pmt_context is not None}")

        print(f"[api_generate_learning_objectives] Generating for org {organization_id}")
        print(f"[api_generate_learning_objectives] Strategies: {len(selected_strategies)}")
        print(f"[api_generate_learning_objectives] PMT customization: {pmt_context is not None}")

        # Generate learning objectives (Week 2 implementation)
        result = generate_complete_learning_objectives(
            org_id=organization_id,
            selected_strategies=selected_strategies,
            pmt_context=pmt_context
        )

        print(f"[api_generate_learning_objectives] Success - Processing time: {result['metadata']['processing_time_seconds']}s")

        return jsonify(result), 200

    except ValueError as e:
        print(f"[api_generate_learning_objectives] Validation error: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'error_type': 'VALIDATION_ERROR'
        }), 400

    except Exception as e:
        print(f"[api_generate_learning_objectives] Unexpected error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': 'An unexpected error occurred during learning objectives generation',
            'error_type': 'INTERNAL_ERROR',
            'details': str(e)
        }), 500


@phase2_learning_bp.route('/phase2/learning-objectives/<int:organization_id>', methods=['GET'])
def api_get_learning_objectives(organization_id):
    """
    Get generated learning objectives for an organization (CACHE ONLY)

    This endpoint only returns cached/previously generated objectives.
    It does NOT trigger LLM generation. Use POST /generate to create new objectives.

    Path Parameters:
        organization_id: Organization ID

    Response (Success - cached data exists):
        {
            "success": true,
            "pathway": "ROLE_BASED" | "TASK_BASED" | "ROLE_BASED_DUAL_TRACK" | "TASK_BASED_DUAL_TRACK",
            "data": {...},
            "metadata": {...}
        }

    Response (No cached data):
        {
            "success": false,
            "error": "No learning objectives generated yet",
            "error_type": "NOT_GENERATED"
        }
    """
    try:
        from app.services.learning_objectives_core import (
            compute_input_hash, get_cached_objectives
        )

        # Validate organization exists
        org = Organization.query.get(organization_id)
        if not org:
            return jsonify({
                'success': False,
                'error': f'Organization {organization_id} not found',
                'error_type': 'ORGANIZATION_NOT_FOUND'
            }), 404

        # Get selected strategies from database
        selected_strategies_db = LearningStrategy.query.filter_by(
            organization_id=organization_id,
            selected=True
        ).order_by(LearningStrategy.priority.asc()).all()

        if not selected_strategies_db:
            return jsonify({
                'success': False,
                'error': 'No learning strategies selected for this organization',
                'error_type': 'NO_STRATEGIES',
                'details': {
                    'message': 'Please select at least one learning strategy in Phase 1 Task 3'
                }
            }), 400

        # Format strategies for hash computation
        selected_strategies = [
            {
                'strategy_id': s.strategy_template_id or s.id,
                'strategy_name': s.strategy_name
            }
            for s in selected_strategies_db
        ]

        # Get PMT context from database (optional)
        pmt_context = None
        pmt_record = OrganizationPMTContext.query.filter_by(organization_id=organization_id).first()
        if pmt_record and pmt_record.is_complete():
            pmt_context = {
                'processes': pmt_record.processes,
                'methods': pmt_record.methods,
                'tools': pmt_record.tools
            }

        # Compute input hash and check cache
        input_hash = compute_input_hash(organization_id, selected_strategies, pmt_context)
        print(f"[api_get_learning_objectives] Checking cache for org {organization_id}")
        print(f"[api_get_learning_objectives] Input hash: {input_hash[:16]}...")

        # Try to get cached result
        cached_result = get_cached_objectives(organization_id, input_hash)

        if cached_result:
            print(f"[api_get_learning_objectives] Cache HIT - returning cached data")
            return jsonify(cached_result), 200
        else:
            # Also check if ANY cache exists (even with different hash)
            # This handles the case where PMT was updated but objectives weren't regenerated
            from models import GeneratedLearningObjectives
            any_cache = GeneratedLearningObjectives.query.filter_by(
                organization_id=organization_id
            ).first()

            if any_cache:
                # Return stale cache with a warning
                print(f"[api_get_learning_objectives] Returning STALE cache (hash mismatch)")
                result = any_cache.objectives_data
                if isinstance(result, str):
                    import json
                    result = json.loads(result)
                result['metadata'] = result.get('metadata', {})
                result['metadata']['from_cache'] = True
                result['metadata']['stale'] = True
                result['metadata']['stale_reason'] = 'Input parameters changed. Click "Generate" to update.'
                return jsonify(result), 200

            # No cache at all - return NOT_GENERATED
            print(f"[api_get_learning_objectives] No cache exists - NOT_GENERATED")
            return jsonify({
                'success': False,
                'error': 'No learning objectives generated yet. Please click "Generate Learning Objectives" to create them.',
                'error_type': 'NOT_GENERATED'
            }), 404

    except ValueError as e:
        print(f"[api_get_learning_objectives] Validation error: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'error_type': 'VALIDATION_ERROR'
        }), 400

    except Exception as e:
        print(f"[api_get_learning_objectives] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': 'An error occurred during learning objectives generation',
            'details': str(e)
        }), 500


# ==============================================================================
# PMT CONTEXT ENDPOINTS
# ==============================================================================

@phase2_learning_bp.route('/phase2/learning-objectives/<int:organization_id>/pmt-context', methods=['GET', 'PATCH'])
def api_pmt_context(organization_id):
    """
    Get or update PMT (Processes, Methods, Tools) context for an organization

    PMT context enables deep customization of learning objectives for 2 specific strategies:
    - "Needs-based project-oriented training"
    - "Continuous support"

    GET Response:
        {
            "organization_id": 28,
            "processes": "Agile development, DevOps deployment",
            "methods": "Scrum, Kanban, TDD",
            "tools": "JIRA, Confluence, Git",
            "industry_specific_context": "Medical device development, ISO 13485",
            "is_complete": true
        }

    PATCH Request Body:
        {
            "processes": "Updated processes",
            "methods": "Updated methods",
            "tools": "Updated tools",
            "industry_specific_context": "Updated industry context"
        }
    """
    try:
        from models import PMTContext

        # Validate organization exists
        org = Organization.query.get(organization_id)
        if not org:
            return jsonify({
                'success': False,
                'error': f'Organization {organization_id} not found'
            }), 404

        if request.method == 'GET':
            # Fetch existing PMT context
            pmt = PMTContext.query.filter_by(organization_id=organization_id).first()

            if not pmt:
                return jsonify({
                    'success': True,
                    'data': {
                        'organization_id': organization_id,
                        'processes': None,
                        'methods': None,
                        'tools': None,
                        'industry': None,
                        'additionalContext': None,
                        'is_complete': False
                    }
                }), 200

            return jsonify({
                'success': True,
                'data': {
                    'organization_id': pmt.organization_id,
                    'processes': pmt.processes,
                    'methods': pmt.methods,
                    'tools': pmt.tools,
                    'industry': pmt.industry,
                    'additionalContext': pmt.additional_context,
                    'is_complete': pmt.is_complete(),
                    'created_at': pmt.created_at.isoformat() if pmt.created_at else None,
                    'updated_at': pmt.updated_at.isoformat() if pmt.updated_at else None
                }
            }), 200

        elif request.method == 'PATCH':
            # Update PMT context
            data = request.get_json()

            if not data:
                return jsonify({
                    'success': False,
                    'error': 'Missing request body'
                }), 400

            # Get or create PMT context
            pmt = PMTContext.query.filter_by(organization_id=organization_id).first()

            if not pmt:
                pmt = PMTContext(organization_id=organization_id)
                db.session.add(pmt)

            # Update fields
            if 'processes' in data:
                pmt.processes = data['processes']
            if 'methods' in data:
                pmt.methods = data['methods']
            if 'tools' in data:
                pmt.tools = data['tools']
            if 'industry' in data:
                pmt.industry = data['industry']
            if 'additionalContext' in data:
                pmt.additional_context = data['additionalContext']

            db.session.commit()

            print(f"[api_pmt_context] Updated PMT context for org {organization_id}")

            return jsonify({
                'success': True,
                'organization_id': pmt.organization_id,
                'processes': pmt.processes,
                'methods': pmt.methods,
                'tools': pmt.tools,
                'industry': pmt.industry,
                'additionalContext': pmt.additional_context,
                'is_complete': pmt.is_complete(),
                'updated_at': pmt.updated_at.isoformat() if pmt.updated_at else None
            }), 200

    except Exception as e:
        print(f"[api_pmt_context] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': 'An error occurred',
            'details': str(e)
        }), 500


# ==============================================================================
# VALIDATION AND PREREQUISITE ENDPOINTS
# ==============================================================================

@phase2_learning_bp.route('/phase2/learning-objectives/<int:organization_id>/validation', methods=['GET'])
def api_get_validation_results(organization_id):
    """
    Get validation results for an organization

    Uses Design v5 validation (Algorithm 2: Mastery Requirements Validation)
    which checks if selected strategies can meet role requirements.

    Response:
        {
            "success": true,
            "organization_id": 28,
            "pathway": "ROLE_BASED" | "TASK_BASED",
            "has_roles": true | false,
            "validation": {
                "status": "OK" | "INADEQUATE",
                "severity": "NONE" | "MEDIUM" | "HIGH",
                "message": "...",
                "affected": [...],
                "recommendations": [...]
            }
        }
    """
    try:
        from app.services.learning_objectives_core import (
            check_if_org_has_roles,
            calculate_combined_targets,
            validate_mastery_requirements
        )

        # Validate organization exists
        org = Organization.query.get(organization_id)
        if not org:
            return jsonify({
                'success': False,
                'error': f'Organization {organization_id} not found'
            }), 404

        # Check if organization has roles
        has_roles = check_if_org_has_roles(organization_id)
        pathway = 'ROLE_BASED' if has_roles else 'TASK_BASED'

        # Get selected strategies from database
        selected_strategies_db = LearningStrategy.query.filter_by(
            organization_id=organization_id,
            selected=True
        ).all()

        if not selected_strategies_db:
            return jsonify({
                'success': False,
                'error': 'No learning strategies selected',
                'error_type': 'NO_STRATEGIES'
            }), 400

        # Format strategies
        selected_strategies = [
            {
                'strategy_id': s.strategy_template_id or s.id,
                'strategy_name': s.strategy_name
            }
            for s in selected_strategies_db
        ]

        # Calculate targets
        targets_result = calculate_combined_targets(selected_strategies)
        main_targets = targets_result['main_targets']

        # Run validation (Algorithm 2)
        validation_result = validate_mastery_requirements(
            organization_id,
            selected_strategies,
            main_targets
        )

        return jsonify({
            'success': True,
            'organization_id': organization_id,
            'pathway': pathway,
            'has_roles': has_roles,
            'validation': validation_result
        }), 200

    except Exception as e:
        print(f"[api_get_validation_results] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': 'An error occurred during validation',
            'details': str(e)
        }), 500


@phase2_learning_bp.route('/phase2/learning-objectives/<int:organization_id>/prerequisites', methods=['GET'])
def api_check_prerequisites(organization_id):
    """
    Check if prerequisites are met for generating learning objectives

    This is a lightweight endpoint for frontend validation before
    enabling the "Generate Objectives" button.

    Response:
        {
            "valid": true,
            "completion_rate": 100.0,
            "pathway": "ROLE_BASED",
            "selected_strategies_count": 2,
            "role_count": 3
        }

    or

        {
            "valid": false,
            "error": "Insufficient assessment data",
            "completion_rate": 45.0,
            "required_rate": 70.0
        }
    """
    try:
        from app.services.pathway_determination import validate_prerequisites

        # Validate organization exists
        org = Organization.query.get(organization_id)
        if not org:
            return jsonify({
                'valid': False,
                'error': f'Organization {organization_id} not found'
            }), 404

        result = validate_prerequisites(organization_id)

        if result.get('valid'):
            return jsonify(result), 200
        else:
            return jsonify(result), 400

    except Exception as e:
        print(f"[api_check_prerequisites] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'valid': False,
            'error': 'An error occurred',
            'details': str(e)
        }), 500


# ==============================================================================
# SETUP AND MANAGEMENT ENDPOINTS
# ==============================================================================

@phase2_learning_bp.route('/phase2/learning-objectives/<int:organization_id>/setup', methods=['POST'])
@jwt_required()
def api_setup_phase2_task3(organization_id):
    """
    Setup Phase 2 Task 3 strategies for an organization

    Automatically creates learning_strategy instances that reference global strategy_template records.
    This endpoint should be called once per organization, typically during organization creation
    or when first enabling Phase 2 Task 3.

    Requirements:
    - User must be authenticated
    - Organization must exist
    - Organization must not already have strategies setup

    Returns:
        200: Setup successful
        400: Setup failed (organization already has strategies)
        404: Organization not found
        500: Server error

    Example request:
        POST /api/phase2/learning-objectives/30/setup

    Example response:
        {
            "success": true,
            "organization_id": 30,
            "organization_name": "New Company",
            "strategies_created": 7,
            "strategies": [
                {"name": "Common basic understanding", "requires_pmt": false},
                ...
            ]
        }
    """
    try:
        print(f"[api_setup_phase2_task3] Setting up Phase 2 Task 3 for org {organization_id}")

        # Verify organization exists
        org = Organization.query.filter_by(id=organization_id).first()
        if not org:
            return jsonify({
                'success': False,
                'error': f'Organization {organization_id} not found'
            }), 404

        # Call setup function (pass db.session for transaction management)
        result = setup_phase2_task3_strategies(organization_id, db_session=db.session)

        if result['success']:
            print(f"[api_setup_phase2_task3] Successfully setup {result['strategies_created']} strategies for org {organization_id}")
            return jsonify(result), 200
        else:
            print(f"[api_setup_phase2_task3] Setup failed for org {organization_id}: {result.get('error')}")
            # If already has strategies, return 400 (client error)
            if 'already has' in result.get('error', ''):
                return jsonify(result), 400
            else:
                return jsonify(result), 500

    except Exception as e:
        print(f"[api_setup_phase2_task3] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': 'An error occurred during setup',
            'details': str(e)
        }), 500


@phase2_learning_bp.route('/phase2/learning-objectives/<int:organization_id>/users', methods=['GET'])
def api_get_assessment_users(organization_id):
    """
    Get detailed list of users and their assessment status

    This endpoint returns a list of all users in the organization with their
    assessment completion status, useful for the Assessment Monitor UI.

    Response:
        {
            "success": true,
            "organization_id": 28,
            "total_users": 9,
            "users_with_assessments": 9,
            "users": [
                {
                    "user_id": 39,
                    "username": "lowmaturity",
                    "email": null,
                    "has_assessment": true,
                    "last_completed": "2025-11-01T17:17:28",
                    "status": "completed"
                },
                ...
            ]
        }
    """
    try:
        from models import User, UserAssessment
        from sqlalchemy import func

        # Validate organization exists
        org = Organization.query.get(organization_id)
        if not org:
            return jsonify({
                'success': False,
                'error': f'Organization {organization_id} not found'
            }), 404

        # Query users with their latest assessment
        users_query = db.session.query(
            User.id,
            User.username,
            User.email,
            func.max(UserAssessment.completed_at).label('last_completed')
        ).outerjoin(
            UserAssessment,
            User.id == UserAssessment.user_id
        ).filter(
            User.organization_id == organization_id
        ).group_by(
            User.id,
            User.username,
            User.email
        ).order_by(
            User.id
        )

        users_data = users_query.all()

        # Format user data
        users_list = []
        users_with_assessments = 0

        for user_id, username, email, last_completed in users_data:
            has_assessment = last_completed is not None
            if has_assessment:
                users_with_assessments += 1

            users_list.append({
                'user_id': user_id,
                'username': username,
                'email': email if email else None,
                'has_assessment': has_assessment,
                'completed_at': last_completed.isoformat() if last_completed else None,
                'status': 'completed' if has_assessment else 'pending'
            })

        return jsonify({
            'success': True,
            'organization_id': organization_id,
            'total_users': len(users_list),
            'users_with_assessments': users_with_assessments,
            'completion_rate': round((users_with_assessments / len(users_list) * 100), 2) if users_list else 0,
            'users': users_list
        }), 200

    except Exception as e:
        print(f"[api_get_assessment_users] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': 'An error occurred',
            'details': str(e)
        }), 500


@phase2_learning_bp.route('/phase2/learning-objectives/<int:organization_id>/add-strategy', methods=['POST'])
def api_add_recommended_strategy(organization_id):
    """
    Add a recommended strategy to organization's selected strategies

    This endpoint allows adding a strategy that was recommended by the validation layer.
    If the strategy requires PMT context, it must be provided in the request.

    Request Body:
        {
            "strategy_name": "Continuous support",
            "pmt_context": {
                "processes": "...",
                "methods": "...",
                "tools": "...",
                "industry_specific_context": "..."
            },
            "regenerate": true
        }

    Response (Success):
        {
            "success": true,
            "message": "Strategy added successfully",
            "strategy": {
                "id": 3,
                "name": "Continuous support",
                "selected": true
            },
            "pmt_required": true,
            "pmt_provided": true,
            "regenerated_objectives": {...}  // Only if regenerate=true
        }

    Response (Error - PMT Missing):
        {
            "success": false,
            "error": "PMT context required",
            "message": "This strategy requires company PMT context for deep customization",
            "pmt_required": true,
            "required_fields": ["processes", "methods", "tools", "industry_specific_context"]
        }
    """
    try:
        from models import PMTContext, LearningStrategy, StrategyTemplate, StrategyTemplateCompetency
        from app.services.learning_objectives_text_generator import check_if_strategy_needs_pmt
        from app.services.pathway_determination import generate_learning_objectives

        data = request.get_json()

        if not data or 'strategy_name' not in data:
            return jsonify({
                'success': False,
                'error': 'Missing strategy_name in request body'
            }), 400

        strategy_name = data['strategy_name']
        pmt_context_data = data.get('pmt_context')
        regenerate = data.get('regenerate', True)  # Default to regenerate

        # Validate organization exists
        org = Organization.query.get(organization_id)
        if not org:
            return jsonify({
                'success': False,
                'error': f'Organization {organization_id} not found'
            }), 404

        print(f"[api_add_recommended_strategy] Adding strategy '{strategy_name}' to org {organization_id}")

        # Find the strategy
        strategy = LearningStrategy.query.filter_by(
            organization_id=organization_id,
            strategy_name=strategy_name
        ).first()

        if not strategy:
            return jsonify({
                'success': False,
                'error': f'Strategy "{strategy_name}" not found for this organization',
                'message': 'Strategy must be created first before adding'
            }), 404

        # Check if strategy already selected
        if strategy.selected:
            return jsonify({
                'success': False,
                'error': 'Strategy already selected',
                'message': f'Strategy "{strategy_name}" is already in selected strategies'
            }), 400

        # Check if strategy requires PMT context
        needs_pmt = check_if_strategy_needs_pmt(strategy_name)

        if needs_pmt:
            # PMT is required for this strategy
            if not pmt_context_data:
                return jsonify({
                    'success': False,
                    'error': 'PMT context required',
                    'message': 'This strategy requires company PMT context for deep customization',
                    'pmt_required': True,
                    'required_fields': ['processes', 'methods', 'tools', 'industry_specific_context']
                }), 400

            # Validate PMT context has required fields
            required_fields = ['processes', 'methods', 'tools', 'industry_specific_context']
            missing_fields = [field for field in required_fields if not pmt_context_data.get(field)]

            if missing_fields:
                return jsonify({
                    'success': False,
                    'error': 'Incomplete PMT context',
                    'message': f'Missing required PMT fields: {", ".join(missing_fields)}',
                    'pmt_required': True,
                    'missing_fields': missing_fields
                }), 400

            # Save or update PMT context
            pmt = PMTContext.query.filter_by(organization_id=organization_id).first()

            if not pmt:
                pmt = PMTContext(organization_id=organization_id)
                db.session.add(pmt)

            pmt.processes = pmt_context_data['processes']
            pmt.methods = pmt_context_data['methods']
            pmt.tools = pmt_context_data['tools']
            pmt.industry_specific_context = pmt_context_data['industry_specific_context']

            db.session.commit()

            print(f"[api_add_recommended_strategy] PMT context updated for org {organization_id}")

        # Mark strategy as selected
        strategy.selected = True

        # Set priority (highest + 1)
        max_priority = db.session.query(
            db.func.max(LearningStrategy.priority)
        ).filter_by(
            organization_id=organization_id,
            selected=True
        ).scalar()

        strategy.priority = (max_priority or 0) + 1

        db.session.commit()

        print(f"[api_add_recommended_strategy] Strategy '{strategy_name}' marked as selected with priority {strategy.priority}")

        # Prepare response
        response = {
            'success': True,
            'message': f'Strategy "{strategy_name}" added successfully',
            'strategy': {
                'id': strategy.id,
                'name': strategy.strategy_name,
                'description': strategy.strategy_description,
                'selected': strategy.selected,
                'priority': strategy.priority
            },
            'pmt_required': needs_pmt,
            'pmt_provided': needs_pmt and bool(pmt_context_data)
        }

        # Regenerate objectives if requested
        if regenerate:
            print(f"[api_add_recommended_strategy] Regenerating objectives with new strategy")
            objectives_result = generate_learning_objectives(organization_id)

            if objectives_result.get('success'):
                response['regenerated_objectives'] = objectives_result
                print(f"[api_add_recommended_strategy] Objectives regenerated successfully")
            else:
                print(f"[api_add_recommended_strategy] Warning: Regeneration failed: {objectives_result.get('error')}")
                response['regeneration_warning'] = objectives_result.get('error')

        return jsonify(response), 200

    except Exception as e:
        print(f"[api_add_recommended_strategy] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': 'An error occurred',
            'details': str(e)
        }), 500


# ==============================================================================
# EXPORT ENDPOINTS
# ==============================================================================

@phase2_learning_bp.route('/phase2/learning-objectives/<int:organization_id>/export', methods=['GET'])
def api_export_learning_objectives(organization_id):
    """
    Export learning objectives in various formats

    Query Parameters:
        format: 'json' | 'excel' | 'pdf' (required)
        strategy: Filter by specific strategy name (optional)
        include_validation: Include validation results (default: true)

    Examples:
        /api/phase2/learning-objectives/28/export?format=json
        /api/phase2/learning-objectives/28/export?format=excel&strategy=Foundation Workshop
        /api/phase2/learning-objectives/28/export?format=pdf&include_validation=true

    Response: File download with appropriate Content-Type
    """
    try:
        from flask import send_file, make_response
        from models import GeneratedLearningObjectives
        import io
        import json as json_lib
        from datetime import datetime

        # Get query parameters
        export_format = request.args.get('format', '').lower()
        strategy_filter = request.args.get('strategy')
        include_validation = request.args.get('include_validation', 'true').lower() == 'true'

        # Validate format
        if export_format not in ['json', 'excel', 'pdf']:
            return jsonify({
                'success': False,
                'error': 'Invalid format',
                'message': 'Format must be one of: json, excel, pdf',
                'valid_formats': ['json', 'excel', 'pdf']
            }), 400

        # Validate organization exists
        org = Organization.query.get(organization_id)
        if not org:
            return jsonify({
                'success': False,
                'error': f'Organization {organization_id} not found'
            }), 404

        print(f"[api_export_learning_objectives] Exporting objectives for org {organization_id} in {export_format} format")

        # Get cached learning objectives (prefer cached data for consistency with UI)
        cached = GeneratedLearningObjectives.query.filter_by(organization_id=organization_id).first()

        if not cached:
            return jsonify({
                'success': False,
                'error': 'No learning objectives generated yet',
                'message': 'Please generate learning objectives first from the LO dashboard'
            }), 400

        # Parse cached data
        result = cached.objectives_data
        if isinstance(result, str):
            result = json_lib.loads(result)

        if not result.get('success', True):
            return jsonify({
                'success': False,
                'error': 'Cached objectives contain errors',
                'details': result.get('error')
            }), 400

        # Filter by strategy if specified
        objectives_data = result.copy()
        if strategy_filter:
            filtered_objectives = {}
            for strategy_id, strategy_data in result.get('learning_objectives_by_strategy', {}).items():
                if strategy_data.get('strategy_name') == strategy_filter:
                    filtered_objectives[strategy_id] = strategy_data

            if not filtered_objectives:
                return jsonify({
                    'success': False,
                    'error': 'Strategy not found',
                    'message': f'No objectives found for strategy "{strategy_filter}"'
                }), 404

            objectives_data['learning_objectives_by_strategy'] = filtered_objectives

        # Remove validation results if not requested
        if not include_validation:
            objectives_data.pop('strategy_validation', None)
            objectives_data.pop('strategic_decisions', None)
            objectives_data.pop('cross_strategy_coverage', None)

        # Export based on format
        if export_format == 'json':
            return export_json(objectives_data, org.organization_name)

        elif export_format == 'excel':
            return export_excel(objectives_data, org.organization_name)

        elif export_format == 'pdf':
            return export_pdf(objectives_data, org.organization_name)

    except Exception as e:
        print(f"[api_export_learning_objectives] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': 'An error occurred during export',
            'details': str(e)
        }), 500


def export_json(data, org_name):
    """Export learning objectives as JSON file"""
    from flask import make_response
    import json as json_lib
    from datetime import datetime

    # Create JSON string with pretty printing
    json_str = json_lib.dumps(data, indent=2, ensure_ascii=False)

    # Create response
    response = make_response(json_str)
    response.headers['Content-Type'] = 'application/json'
    response.headers['Content-Disposition'] = f'attachment; filename="learning_objectives_{org_name.replace(" ", "_")}_{datetime.now().strftime("%Y%m%d")}.json"'

    print(f"[export_json] JSON export created for {org_name}")
    return response


def export_excel(data, org_name):
    """
    Export learning objectives as Excel file matching the Organizational View.

    Single sheet with:
    - Competency rows x Level columns (Knowing, Understanding, Applying)
    - Color coding: Green = Achieved, Yellow = Gap (training required), Gray = Not Targeted
    - LO texts shown as bullet points for Achieved and Training Required
    - Not Targeted cells are gray and show only "Not Targeted" text
    - PMT breakdown shown separately (Process, Method, Tool)
    """
    from flask import send_file
    from datetime import datetime
    import io
    import re

    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return jsonify({
            'success': False,
            'error': 'Excel export not available',
            'message': 'openpyxl library not installed. Install with: pip install openpyxl'
        }), 500

    def format_lo_as_bullets(text):
        """Convert LO text to bullet points by splitting on sentences."""
        if not text:
            return ''
        # Split on periods followed by space or end, but keep sentences meaningful
        sentences = re.split(r'\.(?=\s|$)', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        if len(sentences) <= 1:
            return text  # Single sentence, return as-is
        # Format as bullet points
        return '\n'.join([f"* {s}." if not s.endswith('.') else f"* {s}" for s in sentences])

    def extract_objective_text(lo_data):
        """Extract clean objective text and PMT breakdown from various LO data formats."""
        if not lo_data:
            return '', None

        if isinstance(lo_data, str):
            return lo_data, None

        if isinstance(lo_data, dict):
            pmt_breakdown = None

            # Check for PMT breakdown in the learning_objective object
            if lo_data.get('has_pmt_breakdown') and lo_data.get('pmt_breakdown'):
                pmt_breakdown = lo_data['pmt_breakdown']

            # Check for objective_text field
            if 'objective_text' in lo_data:
                return lo_data['objective_text'], pmt_breakdown

            # Check for direct PMT fields (process, method, tool)
            if 'process' in lo_data or 'method' in lo_data or 'tool' in lo_data:
                # This is the PMT breakdown itself
                return '', {
                    'process': lo_data.get('process', ''),
                    'method': lo_data.get('method', ''),
                    'tool': lo_data.get('tool', '')
                }

            # Fallback: try to get any text-like field
            for key in ['text', 'content', 'description']:
                if key in lo_data:
                    return str(lo_data[key]), pmt_breakdown

        return str(lo_data) if lo_data else '', None

    def format_pmt_breakdown(pmt):
        """Format PMT breakdown with clear labels."""
        if not pmt:
            return ''
        parts = []
        if pmt.get('process'):
            parts.append(f"[PROCESS]\n{pmt['process']}")
        if pmt.get('method'):
            parts.append(f"[METHOD]\n{pmt['method']}")
        if pmt.get('tool'):
            parts.append(f"[TOOL]\n{pmt['tool']}")
        return '\n\n'.join(parts)

    # Color definitions
    HEADER_FILL = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    HEADER_FONT = Font(bold=True, color='FFFFFF')
    ACHIEVED_FILL = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')  # Light green
    GAP_FILL = PatternFill(start_color='FFEB9C', end_color='FFEB9C', fill_type='solid')  # Light yellow/orange
    NOT_TARGETED_FILL = PatternFill(start_color='D9D9D9', end_color='D9D9D9', fill_type='solid')  # Gray
    NOT_TARGETED_FONT = Font(color='808080', italic=True)  # Gray italic text
    TRAINING_EXISTS_FILL = PatternFill(start_color='BDD7EE', end_color='BDD7EE', fill_type='solid')  # Light blue
    TRAINING_EXISTS_FONT = Font(color='2F5496', italic=True)  # Blue italic text
    THIN_BORDER = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Level name mapping (without L1, L2, L4 suffixes)
    LEVEL_NAMES = {
        1: 'Knowing SE',
        2: 'Understanding SE',
        4: 'Applying SE'
    }

    # Create workbook - single sheet only
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.title = 'Learning Objectives'

    # ========== EXTRACT DATA FROM API STRUCTURE ==========
    selected_strategies = data.get('selected_strategies', [])
    if not selected_strategies:
        selected_strategies = data.get('metadata', {}).get('selected_strategies', [])

    all_competencies = {}
    total_gaps = 0
    competencies_with_gaps = set()

    # Check for NEW format (data.main_pyramid)
    main_pyramid = data.get('data', {}).get('main_pyramid', {})
    levels_data = main_pyramid.get('levels', {})

    if levels_data:
        print(f"[export_excel] Using NEW format (data.main_pyramid)")
        for level_str, level_info in levels_data.items():
            level_num = int(level_str)
            if level_num not in [1, 2, 4]:
                continue
            for comp in level_info.get('competencies', []):
                comp_id = comp.get('competency_id')
                comp_name = comp.get('competency_name', f'Competency {comp_id}')
                status = comp.get('status', 'achieved')
                grayed_out = comp.get('grayed_out', False)
                lo_data = comp.get('learning_objective', '')
                target_level = comp.get('target_level', 0)
                current_level = comp.get('current_level', 0)

                if comp_id not in all_competencies:
                    all_competencies[comp_id] = {
                        'name': comp_name,
                        'target_level': target_level,
                        'current_level': current_level,
                        'levels': {}
                    }

                # Get gap_data roles if available
                gap_data = comp.get('gap_data', {})
                roles_data = gap_data.get('roles', {}) if gap_data else {}

                # Extract roles needing this level
                roles_needing = []
                for role_id, role_info in roles_data.items():
                    if isinstance(role_info, dict):
                        level_details = role_info.get('level_details', {}).get(level_num, {})
                        if level_details or level_num in role_info.get('levels_needed', []):
                            roles_needing.append({
                                'role_name': role_info.get('role_name', f'Role {role_id}'),
                                'users_needing': level_details.get('users_needing', role_info.get('users_needing_training', 0)),
                                'total_users': level_details.get('total_users', role_info.get('total_users', 0))
                            })

                all_competencies[comp_id]['levels'][level_num] = {
                    'status': status,
                    'grayed_out': grayed_out,
                    'learning_objective': lo_data,
                    'target_level': target_level,
                    'current_level': current_level,
                    'roles_needing': roles_needing
                }

                if status == 'training_required' and not grayed_out:
                    total_gaps += 1
                    competencies_with_gaps.add(comp_id)
    else:
        print(f"[export_excel] No NEW format data found")

    is_new_format = bool(levels_data)
    print(f"[export_excel] Found {len(all_competencies)} competencies, {total_gaps} gaps")

    # ========== HEADER SECTION ==========
    row = 1
    sheet.merge_cells('A1:D1')
    sheet['A1'] = 'Learning Objectives - Organizational View'
    sheet['A1'].font = Font(size=16, bold=True)
    sheet['A1'].alignment = Alignment(horizontal='center')
    row = 3

    strategy_names = ', '.join([s.get('name', s.get('strategy_name', 'Unknown')) for s in selected_strategies])
    sheet[f'A{row}'] = 'Selected Strategies:'
    sheet[f'A{row}'].font = Font(bold=True)
    sheet[f'B{row}'] = strategy_names if strategy_names else 'None'
    sheet.merge_cells(f'B{row}:D{row}')
    row += 1

    sheet[f'A{row}'] = 'Levels to Advance:'
    sheet[f'A{row}'].font = Font(bold=True)
    sheet[f'B{row}'] = total_gaps
    row += 1

    sheet[f'A{row}'] = 'Competencies with Gap:'
    sheet[f'A{row}'].font = Font(bold=True)
    sheet[f'B{row}'] = len(competencies_with_gaps)
    row += 2

    # Legend
    sheet[f'A{row}'] = 'Legend:'
    sheet[f'A{row}'].font = Font(bold=True)
    row += 1

    sheet[f'A{row}'] = 'Green'
    sheet[f'A{row}'].fill = ACHIEVED_FILL
    sheet[f'B{row}'] = 'Level achieved (no training needed)'
    row += 1

    sheet[f'A{row}'] = 'Yellow'
    sheet[f'A{row}'].fill = GAP_FILL
    sheet[f'B{row}'] = 'Gap exists (training required)'
    row += 1

    sheet[f'A{row}'] = 'Gray'
    sheet[f'A{row}'].fill = NOT_TARGETED_FILL
    sheet[f'A{row}'].font = NOT_TARGETED_FONT
    sheet[f'B{row}'] = 'Not targeted by selected strategies'
    row += 1

    sheet[f'A{row}'] = 'Blue'
    sheet[f'A{row}'].fill = TRAINING_EXISTS_FILL
    sheet[f'A{row}'].font = TRAINING_EXISTS_FONT
    sheet[f'B{row}'] = 'Existing training covers this level (no new training needed)'
    row += 2

    # ========== COMPETENCY TABLE ==========
    matrix_start_row = row
    sorted_comp_ids = sorted(all_competencies.keys())

    # Headers: Competency | Knowing SE | Understanding SE | Applying SE (without L1, L2, L4)
    headers = ['Competency', 'Knowing SE', 'Understanding SE', 'Applying SE']
    for col_idx, header in enumerate(headers, 1):
        cell = sheet.cell(row=row, column=col_idx, value=header)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = THIN_BORDER
    row += 1

    for comp_id in sorted_comp_ids:
        comp_data = all_competencies[comp_id]
        comp_name = comp_data['name']

        # Competency name
        name_cell = sheet.cell(row=row, column=1, value=comp_name)
        name_cell.font = Font(bold=True)
        name_cell.alignment = Alignment(vertical='top', wrap_text=True)
        name_cell.border = THIN_BORDER

        if is_new_format:
            # Level columns (1, 2, 4)
            for col_idx, level_num in enumerate([1, 2, 4], 2):
                level_data = comp_data.get('levels', {}).get(level_num, {})
                cell = sheet.cell(row=row, column=col_idx)
                cell.border = THIN_BORDER
                cell.alignment = Alignment(vertical='top', wrap_text=True)

                if not level_data:
                    cell.value = 'Not Targeted'
                    cell.fill = NOT_TARGETED_FILL
                    cell.font = NOT_TARGETED_FONT
                    continue

                status = level_data.get('status', 'achieved')
                grayed_out = level_data.get('grayed_out', False)
                lo_data = level_data.get('learning_objective', '')
                target_level = level_data.get('target_level', 0)
                current_level = level_data.get('current_level', 0)

                # CRITICAL FIX: Must match frontend SimpleCompetencyCard.vue logic
                # 1. If status is already 'not_targeted', keep it
                # 2. If target_level is 0, this level is NOT TARGETED
                # 3. If level_num > target_level (showing higher level than target), it's NOT TARGETED
                # 4. If current_level >= target_level (and target_level > 0), status should be achieved (no gap)
                if status == 'not_targeted':
                    pass  # Keep the status as is
                elif target_level == 0:
                    status = 'not_targeted'
                elif level_num > target_level:
                    # This level column is higher than the competency's target level
                    # So this level is NOT targeted for this competency
                    status = 'not_targeted'
                elif current_level >= target_level:
                    status = 'achieved'

                # Determine the actual status based on backend logic
                # status: 'training_required' | 'achieved' | 'not_targeted'

                if status == 'not_targeted':
                    # NOT TARGETED - gray cell, no LO text
                    cell.value = 'Not Targeted'
                    cell.fill = NOT_TARGETED_FILL
                    cell.font = NOT_TARGETED_FONT

                elif status == 'training_exists':
                    # TRAINING EXISTS - blue cell indicating existing training covers this level
                    cell.fill = TRAINING_EXISTS_FILL
                    cell.font = TRAINING_EXISTS_FONT
                    cell.value = 'Covered by existing training'

                elif status == 'achieved' or (grayed_out and status not in ['training_required', 'training_exists']):
                    # ACHIEVED - green cell with LO text
                    cell.fill = ACHIEVED_FILL

                    # Extract and format LO text
                    lo_text, pmt_breakdown = extract_objective_text(lo_data)

                    content_parts = []

                    if pmt_breakdown:
                        # Show PMT breakdown with clear sections
                        content_parts.append(format_pmt_breakdown(pmt_breakdown))
                    elif lo_text:
                        # Format as bullet points
                        content_parts.append(format_lo_as_bullets(lo_text))

                    cell.value = '\n'.join(content_parts) if content_parts else ''

                elif status == 'training_required' and not grayed_out:
                    # TRAINING REQUIRED - yellow cell with LO text and role info
                    cell.fill = GAP_FILL

                    content_parts = []

                    # Add role/user info for gaps
                    roles = level_data.get('roles_needing', [])
                    if roles:
                        role_strs = [f"{r.get('role_name', '?')} ({r.get('users_needing', 0)}/{r.get('total_users', 0)})" for r in roles]
                        content_parts.append(f"Roles: {', '.join(role_strs)}")
                        content_parts.append('')

                    # Extract and format LO text
                    lo_text, pmt_breakdown = extract_objective_text(lo_data)

                    if pmt_breakdown:
                        # Show PMT breakdown with clear sections
                        content_parts.append(format_pmt_breakdown(pmt_breakdown))
                    elif lo_text:
                        # Format as bullet points
                        content_parts.append(format_lo_as_bullets(lo_text))

                    cell.value = '\n'.join(content_parts) if content_parts else ''

                else:
                    # Fallback for any other case
                    cell.value = '-'
                    cell.fill = NOT_TARGETED_FILL

        row += 1

    # Column widths - increased for better readability
    sheet.column_dimensions['A'].width = 30
    sheet.column_dimensions['B'].width = 65
    sheet.column_dimensions['C'].width = 65
    sheet.column_dimensions['D'].width = 65

    # Set row heights for wrapped text - increased to accommodate full text
    for r in range(matrix_start_row + 1, row):
        sheet.row_dimensions[r].height = 200  # Increased from 120 to 200

    # Save to BytesIO
    excel_file = io.BytesIO()
    wb.save(excel_file)
    excel_file.seek(0)

    filename = f"learning_objectives_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    print(f"[export_excel] Excel export created for {org_name} with {len(all_competencies)} competencies, {total_gaps} gaps")

    return send_file(
        excel_file,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )


def export_pdf(data, org_name):
    """Export learning objectives as PDF file"""
    from flask import make_response
    from datetime import datetime

    # For now, return a simple text-based PDF using reportlab
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        import io
    except ImportError:
        return jsonify({
            'success': False,
            'error': 'PDF export not available',
            'message': 'reportlab library not installed. Install with: pip install reportlab'
        }), 500

    # Create PDF buffer
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#1976D2'))
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=16, textColor=colors.HexColor('#424242'))

    # Title
    story.append(Paragraph(f'Learning Objectives Report', title_style))
    story.append(Spacer(1, 0.3*inch))

    # Organization info
    story.append(Paragraph(f'<b>Organization:</b> {org_name}', styles['Normal']))
    story.append(Paragraph(f'<b>Pathway:</b> {data.get("pathway", "N/A")}', styles['Normal']))
    story.append(Paragraph(f'<b>Completion Rate:</b> {data.get("completion_rate", 0):.1f}%', styles['Normal']))
    story.append(Paragraph(f'<b>Generated:</b> {datetime.now().strftime("%Y-%m-%d %H:%M")}', styles['Normal']))
    story.append(Spacer(1, 0.5*inch))

    # Selected strategies
    story.append(Paragraph('Selected Strategies', heading_style))
    for strategy in data.get('selected_strategies', []):
        story.append(Paragraph(f'• {strategy["name"]} (Priority {strategy["priority"]})', styles['Normal']))
    story.append(Spacer(1, 0.3*inch))

    # Learning objectives per strategy
    for strategy_id, strategy_data in data.get('learning_objectives_by_strategy', {}).items():
        story.append(PageBreak())

        strategy_name = strategy_data.get('strategy_name', f'Strategy {strategy_id}')
        story.append(Paragraph(strategy_name, heading_style))
        story.append(Spacer(1, 0.2*inch))

        # Summary
        summary = strategy_data.get('summary', {})
        story.append(Paragraph(f'<b>Summary:</b>', styles['Normal']))
        story.append(Paragraph(f'• Training Required: {summary.get("competencies_requiring_training", 0)} competencies', styles['Normal']))
        story.append(Paragraph(f'• Targets Achieved: {summary.get("competencies_targets_achieved", 0)} competencies', styles['Normal']))
        story.append(Spacer(1, 0.3*inch))

        # Trainable competencies table
        trainable = strategy_data.get('trainable_competencies', [])
        training_required = [c for c in trainable if c.get('status') == 'training_required']

        if training_required:
            story.append(Paragraph('<b>Learning Objectives:</b>', styles['Normal']))
            story.append(Spacer(1, 0.1*inch))

            for comp in training_required:
                story.append(Paragraph(f'<b>{comp.get("competency_name")}</b> (Gap: {comp.get("current_level", 0)} → {comp.get("target_level", 0)})', styles['Normal']))
                story.append(Paragraph(comp.get('learning_objective', 'N/A'), styles['BodyText']))
                story.append(Spacer(1, 0.2*inch))

    # Build PDF
    doc.build(story)
    buffer.seek(0)

    filename = f"learning_objectives_{org_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"

    print(f"[export_pdf] PDF export created for {org_name}")

    response = make_response(buffer.read())
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'

    return response


# ==============================================================================
# PMT (Process, Method, Tool) Document Extraction Endpoints
# ==============================================================================

@phase2_learning_bp.route('/phase2/extract-pmt-from-document', methods=['POST'])
def extract_pmt_from_document():
    """
    Extract Process, Method, and Tool information from uploaded documents.
    Uses OpenAI to analyze document text and extract structured PMT data.

    Accepts: PDF, DOCX, TXT files
    Returns: Structured PMT data with confidence scores
    """
    try:
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'}), 400

        file = request.files['file']
        organization_id = request.form.get('organization_id')

        if not file or file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400

        if not organization_id:
            return jsonify({'success': False, 'error': 'organization_id is required'}), 400

        # Get file extension
        filename = file.filename.lower()
        current_app.logger.info(f"[PMT Extract] Processing file: {filename}")

        # Extract text based on file type
        extracted_text = None

        if filename.endswith('.txt'):
            extracted_text = file.read().decode('utf-8', errors='ignore')

        elif filename.endswith('.pdf'):
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text())
                extracted_text = '\n'.join(text_parts)
            except Exception as e:
                current_app.logger.error(f"[ERROR] PDF extraction failed: {str(e)}")
                return jsonify({'success': False, 'error': f'Failed to extract text from PDF: {str(e)}'}), 500

        elif filename.endswith('.docx'):
            try:
                import docx
                doc = docx.Document(file)
                text_parts = [para.text for para in doc.paragraphs]
                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_parts.append(cell.text)
                extracted_text = '\n'.join(text_parts)
            except Exception as e:
                current_app.logger.error(f"[ERROR] DOCX extraction failed: {str(e)}")
                return jsonify({'success': False, 'error': f'Failed to extract text from DOCX: {str(e)}'}), 500

        else:
            return jsonify({'success': False, 'error': 'Unsupported file format. Please upload PDF, DOCX, or TXT files.'}), 400

        if not extracted_text or len(extracted_text.strip()) < 50:
            return jsonify({'success': False, 'error': 'Document appears to be empty or too short'}), 400

        # Use OpenAI to extract PMT information
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

        prompt = f"""You are an expert in analyzing Systems Engineering documentation.
Analyze the following document and extract Process, Method, and Tool information.

DEFINITIONS:
- PROCESS: Organizational workflows, procedures, roles/responsibilities, approval gates, review cycles.
  Examples: ISO standards followed, development lifecycle (V-model, Agile), quality procedures, RACI matrices.

- METHOD: Technical techniques and approaches used to perform engineering activities.
  Examples: Requirements analysis methods, modeling techniques (SysML, UML), trade-off analysis, design reviews.

- TOOL: Specific software, platforms, or tools used to support engineering work.
  Examples: DOORS (requirements), JIRA (project management), Catia Magic (modeling), Git (version control).

Analyze the document text and extract relevant information into these categories.
For each item extracted, provide:
1. The name/title of the item
2. A brief description
3. The category (process, method, or tool)
4. Confidence level (high, medium, low)

Document text:
{extracted_text[:12000]}

Return a JSON object with this structure:
{{
  "document_type": "process|method|tool|mixed",
  "document_summary": "Brief summary of what this document describes",
  "processes": [
    {{"name": "...", "description": "...", "confidence": "high|medium|low"}}
  ],
  "methods": [
    {{"name": "...", "description": "...", "confidence": "high|medium|low"}}
  ],
  "tools": [
    {{"name": "...", "description": "...", "confidence": "high|medium|low"}}
  ],
  "suggested_text": {{
    "processes": "Consolidated text description of all processes found",
    "methods": "Consolidated text description of all methods found",
    "tools": "Consolidated text description of all tools found"
  }}
}}

IMPORTANT: In the "suggested_text" fields:
- If you find items in a category, write a concise summary of what was found
- If NOTHING is found for a category, leave it as an EMPTY STRING ""
- Do NOT write messages like "No specific methods were identified" - just use ""

Return ONLY valid JSON, nothing else."""

        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that extracts Process, Method, and Tool information from Systems Engineering documents. Always respond with valid JSON only."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )

            ai_response = response.choices[0].message.content.strip()

            # Remove markdown code blocks if present
            if ai_response.startswith('```'):
                ai_response = ai_response.split('```')[1]
                if ai_response.startswith('json'):
                    ai_response = ai_response[4:]
                ai_response = ai_response.strip()

            # Parse JSON response
            pmt_data = json.loads(ai_response)

            current_app.logger.info(f"[OK] Extracted PMT data from document - Type: {pmt_data.get('document_type')}")
            current_app.logger.info(f"[OK] Found: {len(pmt_data.get('processes', []))} processes, {len(pmt_data.get('methods', []))} methods, {len(pmt_data.get('tools', []))} tools")

            return jsonify({
                'success': True,
                'filename': file.filename,
                'pmt_data': pmt_data
            })

        except json.JSONDecodeError as e:
            current_app.logger.error(f"[ERROR] Failed to parse AI response: {str(e)}")
            current_app.logger.error(f"[ERROR] AI response was: {ai_response}")
            return jsonify({'success': False, 'error': 'Failed to parse PMT information from document'}), 500

    except Exception as e:
        current_app.logger.error(f"[ERROR] PMT document extraction failed: {str(e)}")
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@phase2_learning_bp.route('/phase2/pmt-reference-examples', methods=['GET'])
def get_pmt_reference_examples():
    """
    Get list of available PMT reference example files.
    Returns metadata about example files that users can view/download.
    """
    try:
        import os
        # Use Docker-safe path helper instead of current_app.root_path
        examples_dir = _get_pmt_examples_dir()

        examples = []

        if os.path.exists(examples_dir):
            for filename in os.listdir(examples_dir):
                if filename.endswith('.txt') and filename.startswith('EXAMPLE_'):
                    filepath = os.path.join(examples_dir, filename)

                    # Determine type from filename
                    if 'PROCESS' in filename:
                        pmt_type = 'process'
                    elif 'METHOD' in filename:
                        pmt_type = 'method'
                    elif 'TOOL' in filename:
                        pmt_type = 'tool'
                    else:
                        pmt_type = 'unknown'

                    # Get file size
                    file_size = os.path.getsize(filepath)

                    # Create friendly name
                    friendly_name = filename.replace('EXAMPLE_', '').replace('_', ' ').replace('.txt', '')

                    examples.append({
                        'filename': filename,
                        'name': friendly_name,
                        'type': pmt_type,
                        'size': file_size,
                        'description': f'Example {pmt_type.upper()} document for reference'
                    })

        return jsonify({
            'success': True,
            'examples': examples
        })

    except Exception as e:
        current_app.logger.error(f"[ERROR] Failed to get PMT examples: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@phase2_learning_bp.route('/phase2/pmt-reference-examples/<filename>', methods=['GET'])
def download_pmt_reference_example(filename):
    """
    Download a specific PMT reference example file.
    """
    try:
        import os
        from flask import send_file

        # Security: only allow specific example files
        if not filename.startswith('EXAMPLE_') or not filename.endswith('.txt'):
            return jsonify({'success': False, 'error': 'Invalid filename'}), 400

        # Prevent directory traversal
        if '..' in filename or '/' in filename or '\\' in filename:
            return jsonify({'success': False, 'error': 'Invalid filename'}), 400

        # Use Docker-safe path helper instead of current_app.root_path
        examples_dir = _get_pmt_examples_dir()
        filepath = os.path.join(examples_dir, filename)

        if not os.path.exists(filepath):
            return jsonify({'success': False, 'error': 'File not found'}), 404

        return send_file(
            filepath,
            mimetype='text/plain',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        current_app.logger.error(f"[ERROR] Failed to download PMT example: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ==============================================================================
# EXISTING TRAINING OFFERS ENDPOINTS
# Feature: "Check and Integrate Existing Offers" (Ulf's request - 11.12.2025)
# ==============================================================================

@phase2_learning_bp.route('/phase2/existing-trainings/<int:organization_id>', methods=['GET'])
def api_get_existing_trainings(organization_id):
    """
    Get list of competencies marked as having existing training in the organization.

    Returns both the list of excluded competency IDs and the full details,
    plus all competencies for the selection UI.

    Response:
        {
            "success": true,
            "data": {
                "existing_training_competencies": [1, 5, 7],
                "existing_trainings_detail": [...],
                "all_competencies": [...]
            }
        }
    """
    try:
        # Validate organization exists
        org = Organization.query.get(organization_id)
        if not org:
            return jsonify({
                'success': False,
                'error': f'Organization {organization_id} not found'
            }), 404

        # Get existing training entries for this organization
        existing = OrganizationExistingTraining.query.filter_by(
            organization_id=organization_id
        ).all()

        # Get all competencies for selection UI
        all_competencies = Competency.query.order_by(Competency.id).all()

        return jsonify({
            'success': True,
            'data': {
                'existing_training_competencies': [e.competency_id for e in existing],
                'existing_trainings_detail': [e.to_dict() for e in existing],
                'all_competencies': [c.to_dict() for c in all_competencies]
            }
        }), 200

    except Exception as e:
        print(f"[api_get_existing_trainings] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': 'An error occurred',
            'details': str(e)
        }), 500


@phase2_learning_bp.route('/phase2/existing-trainings/<int:organization_id>', methods=['PUT'])
def api_update_existing_trainings(organization_id):
    """
    Update the list of competencies with existing training.

    Supports two formats:

    NEW FORMAT (with level differentiation - 2-step process):
        {
            "trainings": [
                {"competency_id": 1, "covered_levels": [1, 2]},
                {"competency_id": 5, "covered_levels": [1, 2, 4]},
                {"competency_id": 7, "covered_levels": [1]}
            ],
            "username": "admin"
        }

    LEGACY FORMAT (backward compatible - all levels excluded):
        {
            "competency_ids": [1, 5, 7],
            "username": "admin"
        }

    Response:
        {
            "success": true,
            "message": "Updated existing trainings: 3 competencies marked",
            "trainings": [...]
        }

    Side Effects:
        - Invalidates the LO cache for this organization (requires regeneration)
    """
    try:
        # Validate organization exists
        org = Organization.query.get(organization_id)
        if not org:
            return jsonify({
                'success': False,
                'error': f'Organization {organization_id} not found'
            }), 404

        data = request.get_json() or {}
        username = data.get('username', 'system')

        # Determine format: new (trainings) or legacy (competency_ids)
        trainings_data = data.get('trainings', None)
        competency_ids = data.get('competency_ids', None)

        # Build list of entries to create
        entries_to_create = []

        if trainings_data is not None:
            # NEW FORMAT: trainings with level differentiation
            for training in trainings_data:
                comp_id = training.get('competency_id')
                covered_levels = training.get('covered_levels', [1, 2, 4])

                # Validate competency exists
                if not Competency.query.get(comp_id):
                    return jsonify({
                        'success': False,
                        'error': f'Invalid competency ID: {comp_id}'
                    }), 400

                # Validate levels
                valid_levels = {1, 2, 4}
                if not all(lvl in valid_levels for lvl in covered_levels):
                    return jsonify({
                        'success': False,
                        'error': f'Invalid levels for competency {comp_id}. Valid levels: 1, 2, 4'
                    }), 400

                entries_to_create.append({
                    'competency_id': comp_id,
                    'covered_levels': json.dumps(sorted(covered_levels))
                })

        elif competency_ids is not None:
            # LEGACY FORMAT: just competency IDs (all levels)
            if competency_ids:
                valid_comps = Competency.query.filter(Competency.id.in_(competency_ids)).all()
                valid_ids = {c.id for c in valid_comps}
                invalid_ids = set(competency_ids) - valid_ids
                if invalid_ids:
                    return jsonify({
                        'success': False,
                        'error': f'Invalid competency IDs: {list(invalid_ids)}'
                    }), 400

                for comp_id in competency_ids:
                    entries_to_create.append({
                        'competency_id': comp_id,
                        'covered_levels': '[1, 2, 4]'  # All levels
                    })

        # Clear existing entries for this organization
        OrganizationExistingTraining.query.filter_by(
            organization_id=organization_id
        ).delete()

        # Add new entries
        for entry_data in entries_to_create:
            entry = OrganizationExistingTraining(
                organization_id=organization_id,
                competency_id=entry_data['competency_id'],
                covered_levels=entry_data['covered_levels'],
                created_by=username
            )
            db.session.add(entry)

        db.session.commit()

        # Invalidate LO cache since exclusions changed
        try:
            from app.services.learning_objectives_core import invalidate_cache
            invalidate_cache(organization_id)
            print(f"[api_update_existing_trainings] Cache invalidated for org {organization_id}")
        except Exception as cache_error:
            print(f"[api_update_existing_trainings] Cache invalidation warning: {str(cache_error)}")
            # Continue - cache invalidation failure shouldn't block the update

        print(f"[api_update_existing_trainings] Updated existing trainings for org {organization_id}: {len(entries_to_create)} competencies")

        # Fetch created entries to return
        created = OrganizationExistingTraining.query.filter_by(
            organization_id=organization_id
        ).all()

        return jsonify({
            'success': True,
            'message': f'Updated existing trainings: {len(entries_to_create)} competencies marked',
            'trainings': [e.to_dict() for e in created]
        }), 200

    except Exception as e:
        print(f"[api_update_existing_trainings] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': 'An error occurred',
            'details': str(e)
        }), 500
